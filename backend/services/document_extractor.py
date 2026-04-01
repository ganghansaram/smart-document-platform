"""
문서 텍스트 추출 — 공통 모듈

3가지 입력(텍스트 PDF, 스캔 PDF, 텍스트 붙여넣기)을 통합 처리하여
듀얼 포맷(표시용 MD + 연산용 plain text)으로 반환한다.

현재 유사도 모드에서 사용하며, 향후 비교/검증 모드에서도 재사용 가능.
"""
import io
import logging
import re
import tempfile
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF

logger = logging.getLogger(__name__)

# ── 스캔 판정 임계값 ──
MIN_TEXT_CHARS_PER_PAGE = 50

# ── OCR DPI ──
OCR_DPI = 300


def extract_document(
    file_bytes: Optional[bytes] = None,
    ext: str = ".pdf",
    text: Optional[str] = None,
) -> dict:
    """문서를 추출하여 듀얼 포맷으로 반환한다.

    Args:
        file_bytes: 파일 바이트 (PDF 또는 DOCX)
        ext: 확장자 (".pdf" 또는 ".docx")
        text: 텍스트 붙여넣기 (file_bytes와 배타적)

    Returns:
        {
            "markdown": str,       # 표시용 구조화 Markdown
            "plain_text": str,     # 연산용 정규화 텍스트 (마크업 제거)
            "page_count": int | None,
            "is_scanned": bool,    # OCR이 사용되었는지
        }
    """
    if text is not None:
        return _from_text(text)

    if file_bytes is None:
        raise ValueError("file_bytes 또는 text 중 하나는 필수입니다")

    if ext == ".pdf":
        return _from_pdf(file_bytes)
    elif ext == ".docx":
        return _from_docx(file_bytes)
    else:
        raise ValueError(f"지원하지 않는 형식: {ext}")


# ══════════════════════════════════════════
# 입력별 처리
# ══════════════════════════════════════════

def _from_text(text: str) -> dict:
    """텍스트 붙여넣기 → 듀얼 포맷"""
    md = text.strip()
    plain = _md_to_plain(md)
    return {
        "markdown": md,
        "plain_text": plain,
        "page_count": None,
        "is_scanned": False,
    }


def _from_docx(file_bytes: bytes) -> dict:
    """DOCX → 듀얼 포맷"""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))

    md_parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # 스타일 기반 헤딩 감지
        style = (para.style.name or "").lower()
        if "heading 1" in style:
            md_parts.append(f"# {text}")
        elif "heading 2" in style:
            md_parts.append(f"## {text}")
        elif "heading 3" in style:
            md_parts.append(f"### {text}")
        else:
            md_parts.append(text)

    # 테이블 추출
    for table in doc.tables:
        md_parts.append(_docx_table_to_md(table))

    md = "\n\n".join(md_parts)
    plain = _md_to_plain(md)

    return {
        "markdown": md,
        "plain_text": plain,
        "page_count": None,
        "is_scanned": False,
    }


def _from_pdf(file_bytes: bytes) -> dict:
    """PDF → 듀얼 포맷 (스캔 자동 감지 + OCR 폴백)"""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    page_count = len(doc)
    is_scanned = False

    # 페이지별 스캔 여부 판정
    text_pages = []  # PyMuPDF4LLM으로 처리할 페이지 인덱스
    scan_pages = []  # OCR로 처리할 페이지 인덱스
    for page_idx in range(page_count):
        raw_text = doc[page_idx].get_text("text").strip()
        if len(raw_text) >= MIN_TEXT_CHARS_PER_PAGE:
            text_pages.append(page_idx)
        else:
            scan_pages.append(page_idx)
            is_scanned = True

    # ── 텍스트 페이지: PyMuPDF4LLM 일괄 추출 ──
    text_md_map = {}
    if text_pages:
        text_md_map = _extract_text_pages_batch(file_bytes, text_pages, doc)

    # ── 스캔 페이지: OCR ──
    scan_md_map = {}
    for page_idx in scan_pages:
        page_md = _extract_scanned_page(doc[page_idx])
        if page_md.strip():
            scan_md_map[page_idx] = page_md

    doc.close()

    # 페이지 순서대로 합치기
    md_pages = []
    for page_idx in range(page_count):
        if page_idx in text_md_map:
            md_pages.append(text_md_map[page_idx])
        elif page_idx in scan_md_map:
            md_pages.append(scan_md_map[page_idx])

    md = "\n\n---\n\n".join(md_pages)
    plain = _md_to_plain(md)

    return {
        "markdown": md,
        "plain_text": plain,
        "page_count": page_count,
        "is_scanned": is_scanned,
    }


# ══════════════════════════════════════════
# 텍스트 PDF 페이지 추출 (PyMuPDF4LLM)
# ══════════════════════════════════════════

def _extract_text_pages_batch(file_bytes: bytes, page_indices: list, doc) -> dict:
    """텍스트 PDF 페이지들을 PyMuPDF4LLM으로 일괄 추출 → {page_idx: md_text}"""
    result = {}

    try:
        import pymupdf4llm

        # 임시 파일에 원본 PDF 저장 (한 번만)
        tmp_path = Path(tempfile.gettempdir()) / f"verify_extract_{id(file_bytes)}.pdf"
        tmp_path.write_bytes(file_bytes)

        try:
            chunks = pymupdf4llm.to_markdown(
                str(tmp_path),
                pages=page_indices,
                page_chunks=True,
                table_strategy="lines_strict",
            )
            for chunk in (chunks or []):
                meta = chunk.get("metadata", {})
                page_idx = meta.get("page", None)
                text = chunk.get("text", "").strip()
                if page_idx is not None and text:
                    result[page_idx] = text
        finally:
            try:
                tmp_path.unlink(missing_ok=True)
            except OSError:
                pass  # Windows 잠금 — 다음 기동 시 자동 정리

    except Exception as e:
        logger.warning(f"PyMuPDF4LLM 배치 추출 실패, 페이지별 폴백: {e}")

    # PyMuPDF4LLM에서 누락된 페이지는 기본 추출로 폴백
    for page_idx in page_indices:
        if page_idx not in result:
            page_md = _extract_text_page_fallback(doc[page_idx])
            if page_md.strip():
                result[page_idx] = page_md

    return result


def _extract_text_page_fallback(page) -> str:
    """PyMuPDF4LLM 실패 시 기본 PyMuPDF로 추출"""
    parts = []

    # 테이블 영역 수집 (텍스트 블록에서 제외하기 위해)
    table_rects = []
    tables_md = []
    try:
        tables = page.find_tables()
        for table in tables.tables:
            table_rects.append(fitz.Rect(table.bbox))
            tables_md.append(_fitz_table_to_md(table))
    except Exception:
        pass

    # 텍스트 블록 추출 (테이블 영역 제외)
    blocks = page.get_text("blocks")
    for b in blocks:
        if b[6] != 0:  # 이미지 블록 스킵
            continue
        block_rect = fitz.Rect(b[:4])
        # 테이블 영역과 겹치면 스킵
        in_table = any(block_rect.intersects(tr) for tr in table_rects)
        if in_table:
            continue
        text = b[4].strip()
        if text:
            parts.append(text)

    # 테이블을 텍스트 뒤에 추가
    parts.extend(tables_md)

    return "\n\n".join(parts)


# ══════════════════════════════════════════
# 스캔 PDF 페이지 추출 (OCR)
# ══════════════════════════════════════════

def _extract_scanned_page(page) -> str:
    """스캔 페이지를 OCR로 텍스트 추출"""
    parts = []

    # 테이블 영역 감지 + 셀별 OCR
    table_rects = []
    try:
        tables_md, t_rects = _ocr_tables(page)
        table_rects = t_rects
        parts.extend(tables_md)
    except Exception as e:
        logger.warning(f"테이블 OCR 실패: {e}")

    # 나머지 영역 전체 OCR
    try:
        text = _ocr_page(page, exclude_rects=table_rects)
        if text.strip():
            # 테이블 앞에 본문 텍스트를 배치
            parts.insert(0, text.strip())
    except Exception as e:
        logger.warning(f"페이지 OCR 실패: {e}")

    return "\n\n".join(parts)


def _ocr_page(page, exclude_rects: list = None) -> str:
    """페이지 이미지를 OCR하여 텍스트 반환"""
    from rapidocr_onnxruntime import RapidOCR

    ocr = _get_ocr_engine()

    # 페이지를 고해상도 이미지로 변환
    pix = page.get_pixmap(dpi=OCR_DPI)
    img_bytes = pix.tobytes("png")

    result, _ = ocr(img_bytes)
    if not result:
        return ""

    # exclude_rects에 속하는 텍스트 제거
    lines = []
    page_width = page.rect.width
    page_height = page.rect.height
    scale_x = page_width / pix.width
    scale_y = page_height / pix.height

    for item in result:
        bbox, text, confidence = item
        if not text or not text.strip():
            continue

        # OCR bbox를 페이지 좌표로 변환
        if exclude_rects and bbox:
            # bbox: [[x1,y1],[x2,y2],[x3,y3],[x4,y4]]
            cx = sum(p[0] for p in bbox) / 4 * scale_x
            cy = sum(p[1] for p in bbox) / 4 * scale_y
            center = fitz.Point(cx, cy)
            in_excluded = any(r.contains(center) for r in exclude_rects)
            if in_excluded:
                continue

        lines.append(text.strip())

    return "\n".join(lines)


def _ocr_tables(page) -> tuple:
    """페이지에서 테이블을 감지하고 셀별 OCR → Markdown 테이블 반환"""
    ocr = _get_ocr_engine()
    tables_md = []
    table_rects = []

    try:
        tables = page.find_tables()
    except Exception:
        return [], []

    for table in tables.tables:
        rect = fitz.Rect(table.bbox)
        table_rects.append(rect)

        # find_tables가 텍스트를 추출할 수 있으면 그대로 사용
        try:
            extracted = table.extract()
            if extracted and any(any(cell for cell in row) for row in extracted):
                tables_md.append(_rows_to_md(extracted))
                continue
        except Exception:
            pass

        # 텍스트 추출 실패 → 테이블 영역 이미지를 OCR
        clip = fitz.Rect(table.bbox)
        pix = page.get_pixmap(dpi=OCR_DPI, clip=clip)
        img_bytes = pix.tobytes("png")
        result, _ = ocr(img_bytes)

        if result:
            text_lines = [item[1].strip() for item in result if item[1].strip()]
            # OCR 결과를 단순 텍스트로 (셀 구조 복원은 어려움)
            tables_md.append("\n".join(text_lines))

    return tables_md, table_rects


# ══════════════════════════════════════════
# 헬퍼 함수
# ══════════════════════════════════════════

# OCR 엔진 싱글턴
_ocr_instance = None

def _get_ocr_engine():
    """RapidOCR 싱글턴 인스턴스"""
    global _ocr_instance
    if _ocr_instance is None:
        from rapidocr_onnxruntime import RapidOCR
        _ocr_instance = RapidOCR()
        logger.info("RapidOCR 엔진 초기화 완료")
    return _ocr_instance


def _md_to_plain(md: str) -> str:
    """Markdown을 연산용 plain text로 변환 (마크업 제거)"""
    text = md

    # 헤딩 마크업 제거 (# ## ### 등)
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)

    # 볼드/이탤릭 제거
    text = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", text)
    text = re.sub(r"_{1,3}(.+?)_{1,3}", r"\1", text)

    # 인라인 코드 제거
    text = re.sub(r"`([^`]+)`", r"\1", text)

    # 링크 → 텍스트만
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)

    # 이미지 제거
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", "", text)

    # GFM 테이블 구분선 제거
    text = re.sub(r"^\|[-:| ]+\|$", "", text, flags=re.MULTILINE)

    # 테이블 파이프 → 공백
    text = re.sub(r"\|", " ", text)

    # 수평선 제거
    text = re.sub(r"^---+$", "", text, flags=re.MULTILINE)

    # 연속 공백/빈 줄 정리
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)

    return text.strip()


def _fitz_table_to_md(table) -> str:
    """PyMuPDF Table 객체를 GFM Markdown 테이블로 변환"""
    try:
        rows = table.extract()
        return _rows_to_md(rows)
    except Exception:
        return ""


def _rows_to_md(rows: list) -> str:
    """2D 문자열 배열을 GFM Markdown 테이블로 변환"""
    if not rows:
        return ""

    def clean(cell):
        return (cell or "").replace("|", "\\|").replace("\n", " ").strip()

    md_rows = []
    for i, row in enumerate(rows):
        cells = [clean(c) for c in row]
        md_rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            md_rows.append("| " + " | ".join(["---"] * len(cells)) + " |")

    return "\n".join(md_rows)


def _docx_table_to_md(table) -> str:
    """python-docx Table 객체를 GFM Markdown 테이블로 변환"""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)
    return _rows_to_md(rows)
