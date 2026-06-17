"""
통일 양식 DOCX 내보내기 서비스 (Plan-60 Phase 3a)

마크다운(SSOT) → HTML → DOCX 파이프라인 + 표지 후처리 주입.

⚠️ 이름 주의: 기존 `export_service.py` 는 Verify/Compare 의 **Excel(.xlsx)** 리포트
생성 서비스로 별개다. 본 모듈은 Plan-60 저작 문서의 **DOCX** 내보내기 전용.

설계 근거 (Phase 1 충실도 PoC, reports/plan-60-phase1-fidelity-poc-2026-06-16.md):
- **2단계 변환** (MD→HTML5 `--mathml` → HTML→DOCX `--reference-doc --number-sections`):
  직접 MD→DOCX 는 병합표(rowspan/colspan)를 셀 단위 문단으로 풀어버려 깨짐.
  HTML 경유 시 병합표·수식(OMML)·이미지·양식이 모두 보존됨.
- **표지**: Pandoc 은 `--reference-doc` 의 본문을 무시하고, HTML 경로는 YAML 제목블록도
  누락하므로, 표지는 python-docx 로 본문 앞에 후처리 주입한다.
"""
import logging
import platform
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn

import config

logger = logging.getLogger(__name__)

# backend/services/ → backend/ → project root
PROJECT_ROOT = Path(__file__).parent.parent.parent
PANDOC_DIR = PROJECT_ROOT / "tools" / "pandoc"
ASSETS_DIR = Path(__file__).parent.parent / "assets"

# 표지/메타에 쓰이는 front matter 키 (통일 양식 사양 §2)
_FM_KEYS = ("title", "subtitle", "author", "date", "doc_number", "classification")

# 표지 색상 (사양 §6 — 통일 블루 계열)
_NAVY = RGBColor(0x1F, 0x38, 0x64)
_RED = RGBColor(0xC0, 0x39, 0x2B)
_GRAY = RGBColor(0x59, 0x59, 0x59)
_COVER_FONT = "맑은 고딕"


class ExportError(Exception):
    """내보내기 변환 실패."""


# ──────────────────────────────────────────────────────────────────────
# 리소스 해석
# ──────────────────────────────────────────────────────────────────────
def resolve_pandoc() -> str:
    """Pandoc 바이너리 경로 해석: config → tools/pandoc/<플랫폼> → PATH."""
    if config.PANDOC_BIN:
        return config.PANDOC_BIN

    if platform.system().lower().startswith("win"):
        candidates = [PANDOC_DIR / "pandoc-windows-amd64.exe", PANDOC_DIR / "pandoc.exe"]
    else:
        candidates = [PANDOC_DIR / "pandoc-linux-amd64", PANDOC_DIR / "pandoc"]
    for c in candidates:
        if c.exists():
            return str(c)

    found = shutil.which("pandoc")
    if found:
        return found
    raise ExportError("Pandoc 바이너리를 찾을 수 없습니다 (tools/pandoc/ 또는 시스템 PATH 확인).")


def resolve_reference_docx() -> Optional[str]:
    """통일 양식 reference.docx 해석: config → data/(관리자 교체) → backend/assets/(기본)."""
    if config.EXPORT_REFERENCE_DOCX:
        p = Path(config.EXPORT_REFERENCE_DOCX)
        if p.exists():
            return str(p)
        logger.warning("EXPORT_REFERENCE_DOCX 설정 경로 없음: %s", p)

    data_ref = PROJECT_ROOT / "data" / "reference.docx"
    if data_ref.exists():
        return str(data_ref)

    asset_ref = ASSETS_DIR / "reference.docx"
    if asset_ref.exists():
        return str(asset_ref)

    logger.warning("reference.docx 없음 → Pandoc 기본 스타일로 내보냄")
    return None


# ──────────────────────────────────────────────────────────────────────
# front matter 파싱 (경량 — pyyaml 의존 없이 알려진 키만)
# ──────────────────────────────────────────────────────────────────────
def parse_front_matter(md_text: str) -> dict:
    """선두 `--- ... ---` 블록에서 알려진 메타 키만 추출."""
    meta: dict = {}
    m = re.match(r"^﻿?---\s*\n(.*?)\n---\s*(?:\n|$)", md_text, re.S)
    if not m:
        return meta
    for line in m.group(1).splitlines():
        mm = re.match(r"\s*([A-Za-z_]+)\s*:\s*(.+?)\s*$", line)
        if mm and mm.group(1) in _FM_KEYS:
            val = mm.group(2).strip().strip('"').strip("'")
            if val:
                meta[mm.group(1)] = val
    return meta


# ──────────────────────────────────────────────────────────────────────
# 표지 주입 (python-docx 후처리)
# ──────────────────────────────────────────────────────────────────────
def _cover_paragraph(doc, first, text, *, size, bold=False, color=None,
                     align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0):
    """표지용 단락을 만들어 본문 첫 요소(first) 앞에 삽입."""
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if text:
        r = p.add_run(text)
        r.font.name = _COVER_FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), _COVER_FONT)
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
    first.addprevious(p._p)
    return p


def inject_cover(docx_path: Path, meta: dict) -> None:
    """본문 docx 맨 앞에 표지 페이지를 주입 (통일 양식 사양 §4.4)."""
    doc = Document(str(docx_path))
    body = doc.element.body
    if len(body) == 0:
        return
    first = body[0]

    if meta.get("classification"):
        _cover_paragraph(doc, first, meta["classification"], size=14, bold=True,
                         color=_RED, before=6, after=2)
    if meta.get("doc_number"):
        _cover_paragraph(doc, first, meta["doc_number"], size=11, color=_NAVY,
                         align=WD_ALIGN_PARAGRAPH.RIGHT, after=120)
    _cover_paragraph(doc, first, meta.get("title", ""), size=28, bold=True,
                     color=_NAVY, before=120, after=10)
    if meta.get("subtitle"):
        _cover_paragraph(doc, first, meta["subtitle"], size=15, color=_GRAY, after=160)
    if meta.get("author"):
        _cover_paragraph(doc, first, meta["author"], size=13, bold=True, after=4)
    if meta.get("date"):
        _cover_paragraph(doc, first, meta["date"], size=12, color=_GRAY, after=0)

    # 표지 끝 페이지 나누기
    pb = doc.add_paragraph()
    pb.add_run().add_break(WD_BREAK.PAGE)
    first.addprevious(pb._p)

    doc.save(str(docx_path))


# ──────────────────────────────────────────────────────────────────────
# 메인 파이프라인
# ──────────────────────────────────────────────────────────────────────
def _run_pandoc(cmd: list, cwd: str) -> None:
    try:
        r = subprocess.run(cmd, cwd=cwd, capture_output=True, text=True,
                           timeout=config.EXPORT_TIMEOUT)
    except subprocess.TimeoutExpired:
        raise ExportError(f"Pandoc 변환 시간 초과 ({config.EXPORT_TIMEOUT}초)")
    except FileNotFoundError:
        raise ExportError("Pandoc 실행 불가 — 바이너리 경로 확인")
    if r.returncode != 0:
        raise ExportError(f"Pandoc 변환 실패: {(r.stderr or r.stdout or '').strip()[:500]}")


def export_markdown_to_docx(md_text: str, *, with_cover: bool = True,
                            resource_path: Optional[str] = None) -> bytes:
    """마크다운 → 통일 양식 DOCX(bytes). 2단계 변환 + 표지 주입.

    Args:
        md_text: 마크다운 원본 (YAML front matter 포함 가능)
        with_cover: front matter 에 title 이 있으면 표지 주입
        resource_path: 이미지 등 상대경로 해석 기준 디렉토리 (서버 내부 경로만 — 클라이언트 입력 금지)
    """
    if not md_text or not md_text.strip():
        raise ExportError("내보낼 마크다운 내용이 비어 있습니다.")

    pandoc = resolve_pandoc()
    reference = resolve_reference_docx()
    meta = parse_front_matter(md_text)

    workdir = Path(tempfile.mkdtemp(prefix="sdp_export_"))
    cwd = str(resource_path) if resource_path else str(workdir)
    try:
        md_path = workdir / "input.md"
        html_path = workdir / "mid.html"
        docx_path = workdir / "out.docx"
        md_path.write_text(md_text, encoding="utf-8")

        # 1단계: MD → HTML5 (수식 MathML 보존)
        _run_pandoc(
            [pandoc, str(md_path), "-f", "markdown", "-t", "html5",
             "--mathml", "-o", str(html_path)],
            cwd=cwd,
        )
        # 2단계: HTML → DOCX (통일 양식 + 섹션 자동번호)
        cmd2 = [pandoc, str(html_path), "-f", "html", "-t", "docx",
                "--number-sections", "-o", str(docx_path)]
        if reference:
            cmd2 += ["--reference-doc", reference]
        if resource_path:
            cmd2 += ["--resource-path", str(resource_path)]
        _run_pandoc(cmd2, cwd=cwd)

        # 3단계: 표지 후처리 주입 (실패해도 본문 내보내기는 유지)
        if with_cover and meta.get("title"):
            try:
                inject_cover(docx_path, meta)
            except Exception as e:
                logger.warning("표지 주입 실패(본문은 정상 내보냄): %s", e)

        return docx_path.read_bytes()
    finally:
        shutil.rmtree(workdir, ignore_errors=True)
