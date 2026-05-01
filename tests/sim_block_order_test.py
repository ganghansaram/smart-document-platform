"""Plan-53 — DOCX 본문/표 원본 순서 보존 검증.

backend/services/docx_utils.py 의 iter_block_items() 가 paragraph 와 table 을
원본 XML 순서대로 yield 하는지 검증한다.

또한 backend/services/document_extractor.py 의 _from_docx() 가 본문 사이에 표가
끼어든 케이스에서 markdown 출력 순서를 보존하는지도 함께 검증.

실행: python tests/sim_block_order_test.py
"""
from __future__ import annotations

import io
import sys
from pathlib import Path

if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8")

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "backend"))

from docx import Document  # noqa: E402
from docx.text.paragraph import Paragraph  # noqa: E402
from docx.table import Table  # noqa: E402

from services.docx_utils import iter_block_items  # noqa: E402


CASES = []


def _build_test_docx(items: list) -> bytes:
    """items: [('p', 'text'), ('t', [['a','b'],['c','d']]), ...] → docx bytes."""
    doc = Document()
    for kind, content in items:
        if kind == 'p':
            doc.add_paragraph(content)
        elif kind == 't':
            table = doc.add_table(rows=len(content), cols=len(content[0]))
            for r, row in enumerate(content):
                for c, val in enumerate(row):
                    table.rows[r].cells[c].text = val
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _items_yielded(file_bytes: bytes) -> list:
    """iter_block_items 가 yield 한 객체들의 (kind, key) 리스트."""
    doc = Document(io.BytesIO(file_bytes))
    out = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            out.append(('p', block.text))
        elif isinstance(block, Table):
            # 첫 셀로 식별
            out.append(('t', block.rows[0].cells[0].text))
    return out


# ─────────────────────────────────────────────────────────────
# Case A — paragraph 만
# ─────────────────────────────────────────────────────────────
def case_only_paragraphs():
    docx = _build_test_docx([
        ('p', 'first'),
        ('p', 'second'),
        ('p', 'third'),
    ])
    items = _items_yielded(docx)
    # 빈 docx 의 시작 1개 단락 무시 — text 필터링
    items = [it for it in items if it[1]]
    assert items == [('p', 'first'), ('p', 'second'), ('p', 'third')], (
        f"paragraph only 순서 어긋남: {items}"
    )


CASES.append(("A. paragraph 만 — 순서 보존", case_only_paragraphs))


# ─────────────────────────────────────────────────────────────
# Case B — table 만
# ─────────────────────────────────────────────────────────────
def case_only_tables():
    docx = _build_test_docx([
        ('t', [['T1-A', 'T1-B'], ['x', 'y']]),
        ('t', [['T2-A', 'T2-B'], ['x', 'y']]),
    ])
    items = _items_yielded(docx)
    tables = [it for it in items if it[0] == 't']
    assert tables == [('t', 'T1-A'), ('t', 'T2-A')], (
        f"table 순서 어긋남: {tables}"
    )


CASES.append(("B. table 만 — 순서 보존", case_only_tables))


# ─────────────────────────────────────────────────────────────
# Case C — paragraph→table→paragraph→table 교차
# 원본 결함: doc.paragraphs 먼저, doc.tables 나중에 → table 이 끝에 모임
# ─────────────────────────────────────────────────────────────
def case_interleaved_p_t():
    docx = _build_test_docx([
        ('p', 'intro'),
        ('t', [['T1-hdr1', 'T1-hdr2'], ['v1', 'v2']]),
        ('p', 'middle'),
        ('t', [['T2-hdr1', 'T2-hdr2'], ['v3', 'v4']]),
        ('p', 'conclusion'),
    ])
    items = _items_yielded(docx)
    items = [it for it in items if it[1]]
    expected = [
        ('p', 'intro'),
        ('t', 'T1-hdr1'),
        ('p', 'middle'),
        ('t', 'T2-hdr1'),
        ('p', 'conclusion'),
    ]
    assert items == expected, (
        f"interleaved 순서 어긋남\n  expected: {expected}\n  got:      {items}"
    )


CASES.append(("C. ★ paragraph↔table 교차 순서 보존", case_interleaved_p_t))


# ─────────────────────────────────────────────────────────────
# Case D — _from_docx markdown 출력 순서 (E2E)
# 본문/표/본문/표/본문 가 markdown 에 그 순서로 등장하는지 검증
# ─────────────────────────────────────────────────────────────
def case_from_docx_e2e_order():
    from services.document_extractor import extract_document  # noqa: E402

    docx = _build_test_docx([
        ('p', 'INTRO_PARAGRAPH'),
        ('t', [['HEADER1_A', 'HEADER1_B'], ['v1', 'v2']]),
        ('p', 'MIDDLE_PARAGRAPH'),
        ('t', [['HEADER2_A', 'HEADER2_B'], ['v3', 'v4']]),
        ('p', 'END_PARAGRAPH'),
    ])
    result = extract_document(file_bytes=docx, ext='.docx')
    md = result.get('markdown', '')

    # 토큰 위치 검사
    positions = {
        'INTRO': md.find('INTRO_PARAGRAPH'),
        'HDR1':  md.find('HEADER1_A'),
        'MIDDLE': md.find('MIDDLE_PARAGRAPH'),
        'HDR2':  md.find('HEADER2_A'),
        'END':   md.find('END_PARAGRAPH'),
    }

    # 모두 발견됐어야 함
    for k, p in positions.items():
        assert p >= 0, f"{k} 토큰 미발견. markdown 일부: {md[:300]}"

    # 순서 검증: INTRO < HDR1 < MIDDLE < HDR2 < END
    order = [positions['INTRO'], positions['HDR1'], positions['MIDDLE'],
             positions['HDR2'], positions['END']]
    assert order == sorted(order), (
        f"markdown 출력 순서 어긋남: {positions}\n"
        f"수정 전 결함: 모든 paragraph 가 모든 table 보다 먼저 나옴 → "
        f"INTRO < MIDDLE < END < HDR1 < HDR2 였음."
    )


CASES.append(("D. ★★ _from_docx markdown 출력 순서 보존 (E2E)", case_from_docx_e2e_order))


# ─────────────────────────────────────────────────────────────
# Case E — 빈 doc
# ─────────────────────────────────────────────────────────────
def case_empty_doc():
    doc = Document()
    out = list(iter_block_items(doc))
    # 빈 문서는 시작 단락 1개만 있을 수 있음
    paragraphs = [b for b in out if isinstance(b, Paragraph)]
    tables = [b for b in out if isinstance(b, Table)]
    assert len(tables) == 0, f"빈 문서에 table {len(tables)}건"
    # paragraph 1개 또는 0개 — 환경에 따라 다름
    assert len(paragraphs) <= 1, f"빈 문서에 paragraph {len(paragraphs)}건"


CASES.append(("E. 빈 doc — 안전 동작", case_empty_doc))


def main() -> int:
    print("Plan-53 — DOCX 본문/표 원본 순서 보존 검증\n")
    fail = 0
    for name, fn in CASES:
        try:
            fn()
            print(f"  PASS  {name}")
        except AssertionError as e:
            print(f"  FAIL  {name}\n        {e}")
            fail += 1
        except Exception as e:  # noqa: BLE001
            print(f"  ERROR {name}\n        {type(e).__name__}: {e}")
            fail += 1

    print()
    if fail == 0:
        print("=" * 60)
        print(f"PASS: {len(CASES)} 케이스 모두 통과 — 본문/표 원본 순서 보장")
        print("=" * 60)
        return 0
    print("=" * 60)
    print(f"FAIL: {fail}/{len(CASES)} 건 실패")
    print("=" * 60)
    return 1


if __name__ == "__main__":
    sys.exit(main())
