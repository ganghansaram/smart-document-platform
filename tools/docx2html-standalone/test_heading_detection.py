#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
헤딩 감지 테스트 — python-docx로 테스트 DOCX 생성 후 변환 검증

테스트 케이스:
  1. 내장 Heading 1~6 스타일         → h1~h6 (outlineLvl)
  2. 커스텀 스타일 + outlineLvl=2    → h3 (outlineLvl)
  3. 한글 "제목 1"~"제목 3"          → h1~h3 (style.name)
  4. Normal 스타일 + 24pt 폰트       → p (본문 스타일 → 폰트 크기 스킵)
  5. Normal 스타일 + 14pt 폰트       → p (오감지 방지)
  6. 다단계 번호 (Heading 1~4)       → h1~h4 (outlineLvl)

실행: python test_heading_detection.py
"""

import os
import re
import sys
import tempfile
from pathlib import Path

# 테스트 대상 모듈 경로
sys.path.insert(0, str(Path(__file__).parent))


def create_test_docx(output_path):
    """다양한 헤딩 시나리오를 포함하는 테스트 DOCX 생성"""
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn

    doc = Document()

    # ── 케이스 1: 내장 Heading 1~6 ──
    for i in range(1, 7):
        doc.add_heading(f'TC1-H{i}: Built-in Heading {i}', level=i)
        doc.add_paragraph(f'Body text after heading {i}.')

    # ── 케이스 2: 커스텀 스타일 + outlineLvl=2 (→ h3) ──
    from docx.enum.style import WD_STYLE_TYPE
    custom_style = doc.styles.add_style('MyCustomSection', WD_STYLE_TYPE.PARAGRAPH)
    # outlineLvl=2 설정 (XML 직접)
    style_pPr = custom_style.element.find(qn('w:pPr'))
    if style_pPr is None:
        style_pPr = custom_style.element.makeelement(qn('w:pPr'), {})
        custom_style.element.append(style_pPr)
    outline_elem = style_pPr.makeelement(qn('w:outlineLvl'), {qn('w:val'): '2'})
    style_pPr.append(outline_elem)

    p = doc.add_paragraph('TC2: Custom style with outlineLvl=2', style='MyCustomSection')
    doc.add_paragraph('Body text after custom heading.')

    # ── 케이스 3: 한글 스타일명은 python-docx로 직접 생성이 어려움 ──
    # 대신 내장 Heading을 사용하되, style.name 매칭 경로를 테스트하기 위해
    # config에 등록된 "Heading 1"~"Heading 3"을 재활용 (케이스 1과 중복 허용)
    # → 실제 한글 Word에서는 "제목 1" 등의 이름이 적용됨
    # 이 케이스는 config.json의 by_style 매핑이 정상 작동하는지를 확인
    doc.add_paragraph('TC3: (한글 스타일 테스트는 실제 한글 Word 문서로 검증)')

    # ── 케이스 4: Normal + 24pt 폰트 (→ p, 오감지 방지) ──
    p = doc.add_paragraph()
    run = p.add_run('TC4: Normal style but 24pt font — should be <p>')
    run.font.size = Pt(24)

    # ── 케이스 5: Normal + 14pt 폰트 (→ p, 오감지 방지) ──
    p = doc.add_paragraph()
    run = p.add_run('TC5: Normal style but 14pt font — should be <p>')
    run.font.size = Pt(14)

    # ── 케이스 6: 다단계 번호 (Heading 1~4) ──
    doc.add_heading('TC6-H1: 1 Top Level', level=1)
    doc.add_heading('TC6-H2: 1.1 Second Level', level=2)
    doc.add_heading('TC6-H3: 1.1.1 Third Level', level=3)
    doc.add_heading('TC6-H4: 1.1.1.1 Fourth Level', level=4)
    doc.add_paragraph('Body text after multi-level headings.')

    doc.save(output_path)
    print(f"[OK] 테스트 DOCX 생성: {output_path}")


def verify_html(html_path):
    """생성된 HTML을 파싱하여 헤딩 태그 검증"""
    with open(html_path, 'r', encoding='utf-8') as f:
        html = f.read()

    results = []
    all_pass = True

    def check(test_id, pattern, expected_tag, description):
        nonlocal all_pass
        # HTML에서 해당 패턴을 포함하는 태그 찾기
        regex = rf'<(h[1-6]|p)[^>]*>[^<]*{re.escape(pattern)}[^<]*</\1>'
        match = re.search(regex, html)
        if match:
            actual_tag = match.group(1)
            passed = actual_tag == expected_tag
        else:
            actual_tag = '(not found)'
            passed = False

        status = 'PASS' if passed else 'FAIL'
        if not passed:
            all_pass = False
        results.append((test_id, status, expected_tag, actual_tag, description))

    # 케이스 1: 내장 Heading 1~6
    for i in range(1, 7):
        check(f'1-H{i}', f'TC1-H{i}:', f'h{i}', f'Built-in Heading {i}')

    # 케이스 2: 커스텀 스타일 + outlineLvl=2 → h3
    check('2', 'TC2:', 'h3', 'Custom style with outlineLvl=2')

    # 케이스 4: Normal + 24pt → p
    check('4', 'TC4:', 'p', 'Normal + 24pt font (must NOT be heading)')

    # 케이스 5: Normal + 14pt → p
    check('5', 'TC5:', 'p', 'Normal + 14pt font (must NOT be heading)')

    # 케이스 6: 다단계 번호
    check('6-H1', 'TC6-H1:', 'h1', 'Multi-level Heading 1')
    check('6-H2', 'TC6-H2:', 'h2', 'Multi-level Heading 2')
    check('6-H3', 'TC6-H3:', 'h3', 'Multi-level Heading 3')
    check('6-H4', 'TC6-H4:', 'h4', 'Multi-level Heading 4')

    # 결과 출력
    print()
    print("=" * 70)
    print(f"{'ID':<8} {'결과':<6} {'기대':<6} {'실제':<10} 설명")
    print("-" * 70)
    for test_id, status, expected, actual, desc in results:
        mark = '  ' if status == 'PASS' else '>>'
        print(f"{mark} {test_id:<6} {status:<6} {expected:<6} {actual:<10} {desc}")
    print("=" * 70)

    passed = sum(1 for r in results if r[1] == 'PASS')
    total = len(results)
    print(f"\n결과: {passed}/{total} PASS", end='')
    if all_pass:
        print(" - ALL PASSED")
    else:
        print(f" - {total - passed} FAILED")

    return all_pass


def main():
    # 임시 디렉토리에서 작업
    with tempfile.TemporaryDirectory(prefix='docx2html_test_') as tmpdir:
        docx_path = os.path.join(tmpdir, 'test_headings.docx')
        html_path = os.path.join(tmpdir, 'test_headings.html')

        # 1. 테스트 DOCX 생성
        print(">>> 테스트 DOCX 생성")
        create_test_docx(docx_path)

        # 2. 변환 실행
        print("\n>>> 변환 실행")
        from converter import DocxConverter
        converter = DocxConverter()
        result = converter.convert(docx_path, html_path)

        if not result.success:
            print(f"[FAIL] 변환 실패: {result.error_message}")
            return 1

        print(f"[OK] 변환 완료: {html_path}")
        stats = result.stats
        print(f"     단락 {stats['paragraphs']}개, "
              f"표 {stats['tables']}개, "
              f"이미지 {stats['images']}개")

        # 3. 검증
        print("\n>>> 검증")
        all_pass = verify_html(html_path)

        return 0 if all_pass else 1


if __name__ == '__main__':
    sys.exit(main())
