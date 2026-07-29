#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
DOCX to HTML Converter - 변환 로직
"""

import json
import os
import re
import sys
import hashlib
from pathlib import Path
from xml.etree import ElementTree as ET

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    raise ImportError("python-docx 패키지가 필요합니다. pip install python-docx 를 실행하세요.")

from utils import (
    get_logger, ensure_dir, get_image_dir, escape_html,
    convert_smart_quotes, ConversionResult, sanitize_filename
)
from omml_to_mathml import OmmlToMathml
from numbering_resolver import NumberingResolver


# ── 표시 캡션 판정 (Plan-73) ──
# 프론트엔드 폴백(js/app.js optimizeContent 2단계)과 의도적으로 동일한 기준.
# 한쪽만 고치면 Explorer(엔진∪JS)와 웹북(엔진 단독)의 화면이 어긋나므로
# 양쪽을 함께 바꾼다.
DISPLAY_CAPTION_RE = re.compile(
    r'^(?:Figure|Table|그림|표|Fig\.)\s*\d', re.IGNORECASE)
DISPLAY_CAPTION_MAX_LEN = 150
# 번호 바로 뒤에 조사가 붙으면 캡션이 아니라 본문 서술이다.
#   "표 1을 보면", "표 2는 …", "그림 3과 같이"  → 본문
#   "표 1 주요 제원", "표 3.1 제원"              → 캡션 (번호 뒤가 공백)
DISPLAY_CAPTION_PARTICLE_RE = re.compile(
    r'^(?:Figure|Table|그림|표|Fig\.)\s*\d+(?:[-.]\d+)*[은는이가을를과와의에도만로으]',
    re.IGNORECASE)


class DocxConverter:
    """Word 문서를 HTML로 변환하는 클래스"""

    def __init__(self, config_path=None):
        """
        Args:
            config_path: 설정 파일 경로 (기본값: ../config.json)
        """
        self.config = self._load_config(config_path)
        self.logger = get_logger()
        self._seq_counters = {}   # SEQ 필드 자동 채번: {category: count}
        self._caption_map = {}    # {caption_id: caption_text}
        self._footnotes = {}      # {id: text}
        self._endnotes = {}       # {id: text}
        self._omml_converter = OmmlToMathml()
        self._shape_warnings = []  # 추출 불가 도형/그리기 경고 목록

    def _load_config(self, config_path):
        """설정 파일 로드 (PyInstaller 번들 호환)"""
        if config_path is None:
            # PyInstaller 번들에선 sys._MEIPASS, 일반 실행엔 __file__ 기준
            base_dir = Path(getattr(sys, '_MEIPASS', Path(__file__).parent))
            config_path = base_dir / "config.json"

        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            return self._default_config()

    def _default_config(self):
        """기본 설정 반환"""
        return {
            "style_mapping": {
                "by_style": {
                    "제목 1": "h1", "제목 2": "h2", "제목 3": "h3",
                    "Heading 1": "h1", "Heading 2": "h2", "Heading 3": "h3",
                    "Title": "h1", "Normal": "p"
                },
                "by_font_size": {
                    "24": "h1", "18": "h2", "14": "h3", "default": "p"
                },
                "priority": "style_first"
            },
            "text_formatting": {
                "bold": "strong", "italic": "em", "underline": "u",
                "strikethrough": "del", "subscript": "sub", "superscript": "sup"
            },
            "special_blocks": {
                "note": ["NOTE", "참고", "비고"],
                "caution": ["CAUTION", "주의"],
                "warning": ["WARNING", "경고"]
            },
            "options": {
                "remove_headers_footers": True,
                "extract_images": True,
                "remove_empty_paragraphs": True,
                "convert_smart_quotes": True
            },
            "output": {
                "fragment_only": True,
                "encoding": "utf-8",
                "indent": True
            }
        }

    def convert(self, input_path, output_path=None, options=None,
                image_dir_name=None, image_prefix=None,
                provenance_adapter=None):
        """
        Word 문서를 HTML로 변환

        Args:
            input_path: 입력 .docx (또는 .docx_1 — DRM 우회용) 파일 경로
            output_path: 출력 .html 파일 경로 (선택)
            options: 변환 옵션 오버라이드 (선택)
            image_dir_name: 이미지 폴더명 오버라이드 (기본: {파일명}_images)
            image_prefix: HTML 내 이미지 경로 접두사 (기본: 상대경로 자동)
            provenance_adapter: HTML provenance meta 에 기록할 전처리 어댑터 이름
                (예: 'word_com', 'libreoffice', 'skip'). 호출자가 디스패처 결과에서
                가져와 전달. None 이면 'unknown'.

        Returns:
            ConversionResult: 변환 결과 정보
        """
        self._image_prefix = image_prefix
        self._provenance_adapter = provenance_adapter
        result = ConversionResult(input_path)
        input_path = Path(input_path)

        # 입력 파일 검증
        if not input_path.exists():
            result.error_message = f"파일을 찾을 수 없습니다: {input_path}"
            return result

        # .docx_1 은 Windows DRM 재적용 우회용 — python-docx 는 ZIP 시그니처로
        # 읽으므로 확장자 무관하게 동작
        if input_path.suffix.lower() not in ('.docx', '.docx_1'):
            result.error_message = f"지원하지 않는 파일 형식입니다: {input_path.suffix}"
            return result

        # 출력 경로 설정
        if output_path is None:
            output_path = input_path.with_suffix('.html')
        output_path = Path(output_path)
        result.output_path = output_path

        # 옵션 병합
        merged_options = {**self.config.get('options', {})}
        if options:
            merged_options.update(options)

        try:
            # Document 로드
            doc = Document(str(input_path))
            self.logger.info(f"문서 로드 완료: {input_path}")

            # 페이지 콘텐츠 폭 계산 (이미지 비율 산출용)
            self._page_content_width = self._get_page_content_width(doc)

            # Phase 4a: numbering.xml 해석기 초기화
            # resolver_mode: 'prefer_cached'(기본) | 'always_polyfill' | 'off'
            resolver_mode = self.config.get('numbering', {}).get(
                'resolver_mode', 'prefer_cached')
            self._numbering_mode = resolver_mode
            if resolver_mode != 'off':
                try:
                    self._numbering_resolver = NumberingResolver(doc)
                except Exception as e:
                    self.logger.warning(f"NumberingResolver 초기화 실패: {e}")
                    self._numbering_resolver = None
            else:
                self._numbering_resolver = None

            # Phase 4b: heading 컨텍스트 (STYLEREF 용) — level(1~9) → 최신 번호 텍스트
            self._heading_context = {i: "" for i in range(1, 10)}

            # 이미지 추출
            image_map = {}
            if merged_options.get('extract_images', True):
                if image_dir_name:
                    image_dir = output_path.parent / image_dir_name
                else:
                    image_dir = get_image_dir(output_path)
                image_map = self._process_images(doc, image_dir, output_path)
                result.stats['images'] = len(image_map)

            # 캡션 카운터/맵 리셋
            self._seq_counters = {}
            self._caption_map = {}
            self._shape_warnings = []

            # 각주/미주 추출
            self._footnotes = self._extract_footnotes(doc)
            self._endnotes = self._extract_endnotes(doc)

            # 리스트 상태 초기화
            self._doc = doc
            self._list_stack = []  # [(numId, ilvl, list_tag)]

            # HTML 생성
            html_parts = []
            first_heading_found = False

            for element in self._iter_block_items(doc):
                if hasattr(element, 'style') and hasattr(element, 'text'):
                    # 리스트 감지
                    list_info = self._get_list_info(element)

                    if list_info is not None:
                        # ── 리스트 문단 처리 ──
                        numId, ilvl, list_tag = list_info
                        target_depth = ilvl + 1
                        transition_html = []
                        opened_new = False

                        # 현재 스택보다 깊이가 얕으면 닫기
                        while len(self._list_stack) > target_depth:
                            _, _, prev_tag = self._list_stack.pop()
                            transition_html.append(f'</li></{prev_tag}>')

                        # 같은 깊이인데 numId가 다르면 기존 닫고 새로 열기
                        if (len(self._list_stack) == target_depth and
                                self._list_stack[-1][0] != numId):
                            _, _, prev_tag = self._list_stack.pop()
                            transition_html.append(f'</li></{prev_tag}>')

                        # 새 리스트 열기 (스택 부족분)
                        while len(self._list_stack) < target_depth:
                            is_target_level = (len(self._list_stack) == target_depth - 1)
                            transition_html.append(f'<{list_tag}>')
                            self._list_stack.append((numId, ilvl, list_tag))
                            opened_new = True
                            if not is_target_level:
                                # 중간 레벨 — 구조적 <li> 래퍼 추가
                                transition_html.append('<li>')

                        # 이전 <li> 닫기 (새로 열지 않은 연속 항목)
                        if not opened_new and self._list_stack:
                            transition_html.append('</li>')

                        if transition_html:
                            html_parts.append(''.join(transition_html))

                        # <li> 내용 생성
                        inner_html = self._get_list_item_html(
                            element, image_map, merged_options)
                        html_parts.append(f'<li>{inner_html}')
                        result.stats['lists'] = result.stats.get('lists', 0) + 1

                    else:
                        # ── 비리스트 → 열린 리스트 모두 닫기 ──
                        close_html = self._close_lists(0)
                        if close_html:
                            html_parts.append(close_html)

                        # 기존 문단 처리
                        html, tag_type = self._process_paragraph(
                            element, image_map, merged_options)

                        if html:
                            if tag_type == 'h1':
                                first_heading_found = True
                            elif not first_heading_found and tag_type != 'h1':
                                result.add_warning("문서가 h1으로 시작하지 않습니다.")

                            html_parts.append(html)
                            result.stats['paragraphs'] += 1

                            if tag_type == 'math':
                                result.stats['equations'] = result.stats.get('equations', 0) + 1

                            if re.match(r'^h[1-6]$', tag_type):
                                result.stats['headings'][tag_type] = result.stats['headings'].get(tag_type, 0) + 1

                elif hasattr(element, 'rows'):
                    # 표 앞에 열린 리스트 닫기
                    close_html = self._close_lists(0)
                    if close_html:
                        html_parts.append(close_html)

                    # 표 처리
                    table_html = self._process_table(element, image_map, merged_options)
                    if table_html:
                        html_parts.append(table_html)
                        result.stats['tables'] += 1

            # 루프 종료 후 남은 리스트 닫기
            close_html = self._close_lists(0)
            if close_html:
                html_parts.append(close_html)

            # 각주/미주 섹션 추가
            all_notes = {**self._footnotes, **self._endnotes}
            if all_notes:
                fn_parts = ['<section class="footnotes"><hr><ol>']
                for fn_id, text in all_notes.items():
                    fn_parts.append(
                        f'<li id="fn-{fn_id}">{escape_html(text)} '
                        f'<a href="#fnref-{fn_id}">\u21a9</a></li>')
                fn_parts.append('</ol></section>')
                html_parts.append('\n'.join(fn_parts))

            # HTML 결합
            indent = self.config.get('output', {}).get('indent', True)
            if indent:
                html_content = '\n\n'.join(html_parts)
            else:
                html_content = ''.join(html_parts)

            # 캡션 참조 하이퍼링크 생성
            html_content = self._linkify_references(html_content)

            # Provenance meta 삽입 (fragment 모드는 주석, 전체 HTML 모드는 <meta>)
            html_content = self._embed_provenance(html_content, merged_options)

            # h1으로 시작하지 않는 경우 경고
            if not first_heading_found:
                result.add_warning("문서에 h1 제목이 없습니다.")

            # 추출 불가 도형/그리기 경고
            if self._shape_warnings:
                n = len(self._shape_warnings)
                result.stats['unextractable_shapes'] = n
                result.add_warning(
                    f"Word 도형/그리기 {n}개가 이미지로 변환되지 않았습니다. "
                    "Word에서 해당 도형 선택 → 복사 → 선택하여 붙여넣기 → "
                    "\"그림(PNG)\"으로 변환 후 재변환하세요."
                )

            # 파일 저장
            ensure_dir(output_path.parent)
            encoding = self.config.get('output', {}).get('encoding', 'utf-8')

            with open(output_path, 'w', encoding=encoding) as f:
                f.write(html_content)

            self.logger.info(f"변환 완료: {output_path}")
            result.success = True

        except Exception as e:
            result.error_message = str(e)
            self.logger.error(f"변환 실패: {input_path} - {e}")

        return result

    def _iter_block_items(self, doc):
        """
        문서의 블록 요소(문단, 표)를 순서대로 반복

        Args:
            doc: Document 객체

        Yields:
            Paragraph 또는 Table 객체
        """
        para_idx = 0
        table_idx = 0

        for child in doc.element.body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag == 'p':
                if para_idx < len(doc.paragraphs):
                    yield doc.paragraphs[para_idx]
                    para_idx += 1
            elif tag == 'tbl':
                if table_idx < len(doc.tables):
                    yield doc.tables[table_idx]
                    table_idx += 1

    # 본문 스타일 ID — heading 오감지 방지
    # Word 의 style_id 는 내부 코드 (예: Normal='a') 이므로 style.name 도 함께 검사
    _BODY_STYLE_IDS = frozenset({
        'Normal', 'BodyText', 'BodyText2', 'BodyText3',
        'ListParagraph', 'ListBullet', 'ListNumber',
        'NoSpacing', 'Quote', 'FootnoteText', 'EndnoteText',
    })
    _BODY_STYLE_NAMES = frozenset({
        'Normal', '본문',
        'Body Text', 'Body Text 2', 'Body Text 3',
        'List Paragraph', 'List Bullet', 'List Number',
        'No Spacing', 'Quote', 'Footnote Text', 'Endnote Text',
    })

    def _is_body_style(self, paragraph):
        """단락이 본문 스타일인지 판정. style_id 와 name 양쪽 검사."""
        if not paragraph.style:
            return False
        if paragraph.style.style_id in self._BODY_STYLE_IDS:
            return True
        if paragraph.style.name in self._BODY_STYLE_NAMES:
            return True
        return False

    def _detect_heading_level(self, paragraph):
        """
        문단의 제목 레벨 감지 (4단계 캐스케이드)

        우선순위:
          1. outlineLvl (OOXML 스펙 정의, 로케일 무관)
          2. style_id  (로케일 무관, config 매핑 + 정규식)
          3. style.name (config 매핑 + 정규식, 로케일 의존)
          4. 폰트 크기  (최후 수단, 본문 스타일이면 스킵)

        Returns:
            str: HTML 태그 (h1~h6 또는 p)
        """
        # Tier 1: outlineLvl (OOXML 스펙 — 가장 권위 있는 신호)
        tag = self._check_outline_level(paragraph)
        if tag:
            return tag

        # Tier 2: style_id (로케일 무관)
        tag = self._check_style_id(paragraph)
        if tag:
            return tag

        # Tier 3: style.name (config 매핑 + 정규식)
        tag = self._check_style_name(paragraph)
        if tag:
            return tag

        # Tier 4: 폰트 크기 (최후 수단, 본문 스타일이면 스킵)
        tag = self._check_font_size(paragraph)
        if tag:
            return tag

        return 'p'

    def _get_heading_numbering(self, paragraph):
        """heading 단락의 (numId, ilvl) 추출. 직접 지정 또는 스타일 상속 체인 탐색.

        Returns:
            (num_id: str, ilvl: int) 또는 None
        """
        # 1. 단락 pPr 에 직접 지정
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                numId_el = numPr.find(qn('w:numId'))
                ilvl_el = numPr.find(qn('w:ilvl'))
                if numId_el is not None:
                    num_id = numId_el.get(qn('w:val'))
                    ilvl = int(ilvl_el.get(qn('w:val'))) if ilvl_el is not None else 0
                    if num_id and num_id != '0':
                        return (num_id, ilvl)

        # 2. 스타일 상속 체인 (Heading 1 스타일에 numPr 정의된 경우)
        style = paragraph.style
        depth = 0
        while style is not None and depth < 10:
            style_elem = style.element
            if style_elem is not None:
                style_pPr = style_elem.find(qn('w:pPr'))
                if style_pPr is not None:
                    numPr = style_pPr.find(qn('w:numPr'))
                    if numPr is not None:
                        numId_el = numPr.find(qn('w:numId'))
                        ilvl_el = numPr.find(qn('w:ilvl'))
                        if numId_el is not None:
                            num_id = numId_el.get(qn('w:val'))
                            ilvl = int(ilvl_el.get(qn('w:val'))) if ilvl_el is not None else 0
                            if num_id and num_id != '0':
                                return (num_id, ilvl)
            style = style.base_style
            depth += 1

        return None

    def _check_outline_level(self, paragraph):
        """Tier 1: outlineLvl — 단락 XML 직접 조회 + 스타일 상속 체인.

        본문 스타일(_BODY_STYLE_IDS/_BODY_STYLE_NAMES)에 속하는 단락은
        outlineLvl 이 지정되어 있어도 heading 으로 인식하지 않는다 — 일부
        Word 템플릿이 TOC 용으로 본문 단락에 outlineLvl 을 지정하는 경우
        오감지 방지.
        """
        # 본문 스타일이면 outlineLvl 존재 여부와 무관하게 heading 아님
        if self._is_body_style(paragraph):
            return None

        # 단락 자체의 pPr에서 outlineLvl 확인
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is not None:
            outline = pPr.find(qn('w:outlineLvl'))
            if outline is not None:
                val = int(outline.get(qn('w:val')))
                if val <= 8:
                    return f'h{min(val + 1, 6)}'
                return None  # val=9 → 본문

        # 스타일 정의의 pPr에서 outlineLvl 확인 (상속 체인 탐색)
        style = paragraph.style
        depth = 0
        while style is not None and depth < 10:
            style_elem = style.element
            if style_elem is not None:
                style_pPr = style_elem.find(qn('w:pPr'))
                if style_pPr is not None:
                    outline = style_pPr.find(qn('w:outlineLvl'))
                    if outline is not None:
                        val = int(outline.get(qn('w:val')))
                        if val <= 8:
                            return f'h{min(val + 1, 6)}'
                        return None  # val=9 → 본문
            style = style.base_style
            depth += 1

        return None

    def _check_style_id(self, paragraph):
        """Tier 2: style_id — 로케일 무관 매핑"""
        if not paragraph.style or not paragraph.style.style_id:
            return None

        style_id = paragraph.style.style_id
        style_mapping = self.config.get('style_mapping', {})

        # config의 by_style_id 매핑 조회
        by_style_id = style_mapping.get('by_style_id', {})
        tag = by_style_id.get(style_id)
        if tag:
            return tag

        # 정규식 폴백: "Heading1" ~ "Heading9"
        match = re.match(r'^Heading(\d+)$', style_id)
        if match:
            return f'h{min(int(match.group(1)), 6)}'

        return None

    def _check_style_name(self, paragraph):
        """Tier 3: style.name — config 매핑 + 정규식"""
        if not paragraph.style or not paragraph.style.name:
            return None

        style_name = paragraph.style.name
        style_mapping = self.config.get('style_mapping', {})
        by_style = style_mapping.get('by_style', {})

        tag = by_style.get(style_name)
        if tag and tag != 'p':
            return tag

        # 정규식 폴백: "Heading N", "제목 N" 등
        match = re.match(r'^(?:Heading|제목)\s*(\d+)$', style_name, re.IGNORECASE)
        if match:
            return f'h{min(int(match.group(1)), 6)}'

        return None

    def _check_font_size(self, paragraph):
        """Tier 4: 폰트 크기 — 최후 수단. 본문 스타일이면 스킵."""
        # 본문 스타일이면 폰트 크기와 무관하게 헤딩 아님
        if self._is_body_style(paragraph):
            return None

        style_mapping = self.config.get('style_mapping', {})
        by_font_size = style_mapping.get('by_font_size', {})
        font_size = self._get_paragraph_font_size(paragraph)
        if font_size:
            font_size_str = str(int(font_size))
            tag = by_font_size.get(font_size_str)
            if tag and tag != 'p':
                return tag

        return None

    def _get_paragraph_font_size(self, paragraph):
        """
        문단의 폰트 크기 추출 (포인트 단위)

        Args:
            paragraph: Paragraph 객체

        Returns:
            float or None: 폰트 크기
        """
        # 첫 번째 run의 폰트 크기 확인
        for run in paragraph.runs:
            if run.font.size:
                return run.font.size.pt

        # 스타일에서 폰트 크기 확인
        if paragraph.style and paragraph.style.font and paragraph.style.font.size:
            return paragraph.style.font.size.pt

        return None

    def _process_paragraph(self, paragraph, image_map, options):
        """
        문단을 HTML로 변환

        Args:
            paragraph: Paragraph 객체
            image_map: 이미지 매핑 딕셔너리
            options: 변환 옵션

        Returns:
            tuple: (HTML 문자열, 태그 타입)
        """
        # TOC(목차) 문단 스킵 — 우측 On this page 패널이 대체
        if paragraph.style and paragraph.style.name:
            sn = paragraph.style.name.lower()
            if sn.startswith('toc') or sn == 'table of figures':
                return None, None

        # SEQ 필드 해석 (빈 캐시 값 자동 채번) + STYLEREF 해석 (Phase 4b)
        self._resolve_styleref_fields(paragraph)
        self._resolve_seq_fields(paragraph)

        text = paragraph.text.strip()
        has_math = self._has_math(paragraph)

        # 빈 문단 처리
        if not text and not self._has_images(paragraph) and not has_math:
            if options.get('remove_empty_paragraphs', True):
                return None, None

        # 특수 블록 감지
        special_block = self._detect_special_block(text)
        if special_block:
            block_type, content = special_block
            return f'<div class="{block_type}">{escape_html(content)}</div>', 'special'

        # 제목 레벨 감지
        tag = self._detect_heading_level(paragraph)

        # Phase 4a: heading 자동번호 polyfill
        # numbering.xml 기반으로 번호 prefix 계산. counter 가 문서 순서대로
        # 누적되므로 이 블록은 "순회 중인 모든 heading" 에 대해 반드시 한 번씩
        # 실행되어야 함 (skip 불가).
        heading_number_prefix = ""
        if tag and tag.startswith('h') and self._numbering_resolver is not None:
            num_info = self._get_heading_numbering(paragraph)
            if num_info is not None:
                num_id, ilvl = num_info
                try:
                    heading_number_prefix = self._numbering_resolver.format_number(
                        num_id, ilvl) or ""
                except Exception as e:
                    self.logger.debug(f"numbering polyfill 실패: {e}")

        # Phase 4b: heading 컨텍스트 업데이트 (STYLEREF 에서 사용)
        if tag and tag.startswith('h') and hasattr(self, '_heading_context'):
            level = int(tag[1])
            # 우선순위: ① numbering polyfill 결과 → ② 본문 텍스트 앞자리 숫자 →
            # ③ 자체 카운터 fallback (numbering.xml 도 없고 텍스트에도 번호 없는 경우)
            num_text = heading_number_prefix.strip().rstrip('.')
            if not num_text:
                m = re.match(r'^(\d+(?:[\.\-]\d+)*)', text)
                if m:
                    num_text = m.group(1)
            if not num_text:
                # fallback counter — level 별 등장 순서. 상위 level 변경 시 하위 리셋.
                ctr_key = f'_simple_h{level}'
                prev_top_key = f'_simple_prev_parent_{level}'
                if level > 1:
                    parent = self._heading_context.get(level - 1, '')
                    if self._heading_context.get(prev_top_key) != parent:
                        self._heading_context[ctr_key] = 0
                        self._heading_context[prev_top_key] = parent
                self._heading_context[ctr_key] = \
                    self._heading_context.get(ctr_key, 0) + 1
                num_text = str(self._heading_context[ctr_key])
            self._heading_context[level] = num_text
            # 하위 level 은 다음 heading 이 들어올 때 덮어씌워짐 — 현재 level 까지만 업데이트

        # 스타일 속성 구성 (정렬 + 들여쓰기)
        style_parts = []
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            style_parts.append('text-align: center')
        elif paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            style_parts.append('text-align: right')
        elif paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            style_parts.append('text-align: justify')

        # 들여쓰기 확인 (리스트가 아닌 일반 문단)
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is not None:
            ind_elem = pPr.find(qn('w:ind'))
            if ind_elem is not None:
                left_val = ind_elem.get(qn('w:left'))
                if left_val:
                    try:
                        left_twips = int(left_val)
                        if left_twips > 0:
                            margin_px = round(left_twips / 20)
                            style_parts.append(f'margin-left: {margin_px}px')
                    except ValueError:
                        pass

        align_attr = f' style="{"; ".join(style_parts)};"' if style_parts else ''

        # 디스플레이 수식 (m:oMathPara — 독립 블록 수식)
        oMathPara = paragraph._element.find(qn('m:oMathPara'))
        if oMathPara is not None:
            oMath = oMathPara.find(qn('m:oMath'))
            if oMath is not None:
                math_html = self._omml_converter.convert_omath(oMath, display=True)
                return f'<div class="math-display"{align_attr}>{math_html}</div>', 'math'
            return None, None

        # 인라인 서식 처리 (수식·하이퍼링크 혼합 시 자식 요소 직접 순회)
        self._shape_warned_para = False  # 도형 경고 중복 방지
        has_hyperlink = paragraph._element.find(qn('w:hyperlink')) is not None
        if has_math or has_hyperlink:
            inner_html = self._process_paragraph_children(paragraph, image_map, options)
        else:
            inner_html = self._process_runs(paragraph.runs, image_map, options)

        if not inner_html.strip():
            # 이미지만 있는 경우
            images_html = self._extract_inline_images(paragraph, image_map, align_attr)
            if images_html:
                return images_html, 'image'
            # 도형/그리기만 있는 경우 → 플레이스홀더
            if self._has_unextractable_shapes(paragraph):
                self._shape_warnings.append(len(self._shape_warnings) + 1)
                return (
                    '<div class="shape-placeholder">'
                    '⚠ Word 도형/그리기 개체 — 이미지로 변환되지 않았습니다. '
                    'Word에서 해당 도형을 선택 → 복사 → '
                    '선택하여 붙여넣기 → "그림(PNG)"으로 변환 후 재변환하세요.'
                    '</div>',
                    'shape'
                )
            return None, None

        # 캡션 감지 (Plan-73: 참조 캡션 / 표시 캡션 2계층)
        #   참조 캡션 — id + 본문 참조 링크 대상. 판정은 _detect_caption (엄격, 무변경)
        #   표시 캡션 — class 만. 판정은 _is_display_caption (느슨)
        # id 는 참조 캡션에만 부여한다. 표시 캡션까지 id 를 주면 중복 id·본문
        # 오탐 링크가 생겨 semantic_checks·사용자 가이드 계약이 깨진다.
        caption_id = self._detect_caption(text, paragraph)
        id_attr = f' id="{caption_id}"' if caption_id else ''
        is_caption = bool(caption_id) or (
            tag == 'p' and self._is_display_caption(text))
        caption_cls = ' class="caption"' if is_caption else ''

        # Phase 4a: heading 번호 prefix 적용
        # resolver_mode 에 따라:
        #  - prefer_cached (기본): 본문 텍스트에 이미 번호 있으면 스킵
        #  - always_polyfill: 무조건 prefix 주입
        #  - off: 주입 안 함 (사실상 doesn't reach here because resolver_mode='off' skips init)
        if heading_number_prefix and tag and tag.startswith('h'):
            mode = getattr(self, '_numbering_mode', 'prefer_cached')
            should_prepend = True
            if mode == 'prefer_cached':
                # 본문 텍스트가 이미 숫자 prefix 로 시작하면 cached 취급 → 건너뜀
                if re.match(r'^[\d]+(?:[\.\-][\d]+)*[\.\s]', text):
                    should_prepend = False
            if should_prepend:
                inner_html = escape_html(heading_number_prefix) + ' ' + inner_html

        html = f'<{tag}{id_attr}{caption_cls}{align_attr}>{inner_html}</{tag}>'

        # 텍스트와 섞인 도형이 감지된 경우 → 문단 뒤에 블록 플레이스홀더 추가
        if self._shape_warned_para:
            self._shape_warnings.append(len(self._shape_warnings) + 1)
            html += (
                '\n\n<div class="shape-placeholder">'
                '⚠ 이 위치에 Word 도형/그리기 개체가 있으나 이미지로 변환되지 않았습니다. '
                'Word에서 해당 도형을 선택 → 복사 → '
                '선택하여 붙여넣기 → "그림(PNG)"으로 변환 후 재변환하세요.'
                '</div>')

        return html, tag

    def _get_format_tags(self, run):
        """run의 서식을 (여는 태그, 닫는 태그) 문자열로 반환"""
        open_parts, close_parts = [], []
        fmt = self.config.get('text_formatting', {})
        for attr, key in [('bold', 'bold'), ('italic', 'italic'), ('underline', 'underline')]:
            if getattr(run, attr, None):
                t = fmt.get(key, key)
                open_parts.append(f'<{t}>')
                close_parts.insert(0, f'</{t}>')
        if run.font.strike:
            t = fmt.get('strikethrough', 'del')
            open_parts.append(f'<{t}>')
            close_parts.insert(0, f'</{t}>')
        if run.font.subscript:
            t = fmt.get('subscript', 'sub')
            open_parts.append(f'<{t}>')
            close_parts.insert(0, f'</{t}>')
        if run.font.superscript:
            t = fmt.get('superscript', 'sup')
            open_parts.append(f'<{t}>')
            close_parts.insert(0, f'</{t}>')
        # 텍스트 색상 (검정 #000000 은 기본값이므로 스킵 — CSS 우선)
        try:
            if run.font.color and run.font.color.rgb:
                hex_c = str(run.font.color.rgb)
                if hex_c != '000000':
                    open_parts.append(f'<span style="color: #{hex_c}">')
                    close_parts.insert(0, '</span>')
        except (AttributeError, TypeError):
            pass
        # 하이라이트
        try:
            if run.font.highlight_color:
                hl_map = {
                    1: '#000000', 2: '#0000ff', 3: '#00ffff', 4: '#00ff00',
                    5: '#ff00ff', 6: '#ff0000', 7: '#ffff00', 8: '#ffffff',
                    9: '#00008b', 10: '#008080', 11: '#008000', 12: '#800080',
                    13: '#800000', 14: '#808000', 15: '#808080', 16: '#c0c0c0',
                }
                bg = hl_map.get(int(run.font.highlight_color))
                if bg:
                    open_parts.append(f'<mark style="background-color:{bg}">')
                    close_parts.insert(0, '</mark>')
        except (AttributeError, TypeError, ValueError):
            pass
        return ''.join(open_parts), ''.join(close_parts)

    def _process_runs(self, runs, image_map, options):
        """Run 요소들을 HTML로 변환 (내부 XML 자식 직접 순회)"""
        parts = []
        for run in runs:
            open_tags, close_tags = self._get_format_tags(run)
            has_content = False
            for child in run._element:
                ctag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if ctag == 't':
                    text = child.text or ''
                    if options.get('convert_smart_quotes', True):
                        text = convert_smart_quotes(text)
                    text = escape_html(text)
                    parts.append(f'{open_tags}{text}{close_tags}')
                    has_content = True
                elif ctag == 'br':
                    br_type = child.get(qn('w:type'))
                    if br_type == 'page':
                        parts.append('<hr class="page-break">')
                    else:
                        parts.append('<br>')
                    has_content = True
                elif ctag == 'tab':
                    parts.append('&emsp;')
                    has_content = True
                elif ctag == 'footnoteReference':
                    fn_id = child.get(qn('w:id'))
                    if fn_id and fn_id in self._footnotes:
                        parts.append(
                            f'{open_tags}<a href="#fn-{fn_id}" id="fnref-{fn_id}" '
                            f'class="fn-ref">[{fn_id}]</a>{close_tags}')
                        has_content = True
            # 이미지 (텍스트 없는 run에서만)
            if not has_content:
                embed_id, blip_el = self._get_run_image(run)
                if embed_id and embed_id in image_map:
                    parts.append(self._make_img_tag(image_map[embed_id], blip_el))
                elif self._run_has_shape(run):
                    if not getattr(self, '_shape_warned_para', False):
                        self._shape_warned_para = True
        return ''.join(parts)

    def _process_paragraph_children(self, paragraph, image_map, options):
        """문단의 직접 자식을 순회하여 run, 수식, 하이퍼링크, 북마크를 처리."""
        from docx.text.run import Run
        parts = []
        for child in paragraph._element:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'r':
                run = Run(child, paragraph)
                run_html = self._process_runs([run], image_map, options)
                if run_html:
                    parts.append(run_html)
            elif tag == 'oMath':
                parts.append(self._omml_converter.convert_omath(child, display=False))
            elif tag == 'hyperlink':
                r_id = child.get(qn('r:id'))
                anchor = child.get(qn('w:anchor'))
                url = ''
                if r_id:
                    try:
                        rel = paragraph.part.rels[r_id]
                        url = rel.target_ref
                    except (KeyError, AttributeError):
                        pass
                elif anchor:
                    url = f'#{anchor}'
                link_runs = [Run(r, paragraph) for r in child.findall(qn('w:r'))]
                link_text = self._process_runs(link_runs, image_map, options)
                if url:
                    parts.append(f'<a href="{escape_html(url)}">{link_text}</a>')
                else:
                    parts.append(link_text)
            elif tag == 'bookmarkStart':
                bm_name = child.get(qn('w:name'))
                if bm_name and not bm_name.startswith('_'):
                    parts.append(f'<a id="{escape_html(bm_name)}"></a>')
            # pPr, oMathPara 등은 스킵
        return ''.join(parts)

    # ── 리스트 변환 관련 메서드 ──────────────────────────────────────

    def _get_list_info(self, paragraph):
        """리스트 문단이면 (numId, ilvl, list_tag) 반환, 아니면 None"""
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is None:
            return None
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            return None

        numId_el = numPr.find(qn('w:numId'))
        ilvl_el = numPr.find(qn('w:ilvl'))
        numId = numId_el.get(qn('w:val')) if numId_el is not None else '0'
        ilvl = int(ilvl_el.get(qn('w:val'))) if ilvl_el is not None else 0

        # numId=0은 리스트 없음을 의미
        if numId == '0':
            return None

        list_tag = self._determine_list_tag(numId, ilvl)
        return (numId, ilvl, list_tag)

    def _determine_list_tag(self, numId, ilvl):
        """numId, ilvl로 numbering 정의를 조회하여 'ol' 또는 'ul' 반환"""
        try:
            numbering_part = self._doc.part.numbering_part
            if numbering_part is None:
                return 'ul'
            numbering_elem = numbering_part.element

            # numId → abstractNumId 조회
            abstract_num_id = None
            for num in numbering_elem.findall(qn('w:num')):
                if num.get(qn('w:numId')) == numId:
                    abs_el = num.find(qn('w:abstractNumId'))
                    if abs_el is not None:
                        abstract_num_id = abs_el.get(qn('w:val'))
                    break

            if abstract_num_id is None:
                return 'ul'

            # abstractNum → 해당 레벨의 numFmt 조회
            for abs_num in numbering_elem.findall(qn('w:abstractNum')):
                if abs_num.get(qn('w:abstractNumId')) == abstract_num_id:
                    for lvl in abs_num.findall(qn('w:lvl')):
                        if lvl.get(qn('w:ilvl')) == str(ilvl):
                            num_fmt = lvl.find(qn('w:numFmt'))
                            if num_fmt is not None:
                                fmt = num_fmt.get(qn('w:val'))
                                if fmt == 'bullet':
                                    return 'ul'
                                return 'ol'
                    break
        except Exception:
            pass
        return 'ul'

    def _close_lists(self, target_depth=0):
        """리스트 스택을 target_depth까지 닫기, HTML 조각 반환"""
        parts = []
        while len(self._list_stack) > target_depth:
            _, _, tag = self._list_stack.pop()
            parts.append(f'</li></{tag}>')
        return ''.join(parts)

    def _get_list_item_html(self, paragraph, image_map, options):
        """리스트 항목의 내부 HTML 생성 (래핑 태그 없이)"""
        self._resolve_styleref_fields(paragraph)
        self._resolve_seq_fields(paragraph)
        has_math = self._has_math(paragraph)

        # 디스플레이 수식
        oMathPara = paragraph._element.find(qn('m:oMathPara'))
        if oMathPara is not None:
            oMath = oMathPara.find(qn('m:oMath'))
            if oMath is not None:
                return (f'<div class="math-display">'
                        f'{self._omml_converter.convert_omath(oMath, display=True)}'
                        f'</div>')
            return ''

        # 인라인 수식 또는 하이퍼링크 포함
        has_hyperlink = paragraph._element.find(qn('w:hyperlink')) is not None
        if has_math or has_hyperlink:
            return self._process_paragraph_children(paragraph, image_map, options)

        # 일반 텍스트/서식
        inner = self._process_runs(paragraph.runs, image_map, options)
        if inner.strip():
            return inner

        # 이미지만 있는 경우
        images = []
        try:
            for elem in paragraph._element.iter():
                if 'blip' in elem.tag:
                    embed = elem.get(qn('r:embed'))
                    if embed and embed in image_map:
                        images.append(self._make_img_tag(image_map[embed], elem))
        except Exception:
            pass
        return ''.join(images)

    def _process_cell_content(self, cell, image_map, options):
        """표 셀 내용을 서식 보존하여 변환. (content, align) 튜플 반환."""
        parts = []
        align = None
        for para in cell.paragraphs:
            text = para.text.strip()
            if not text and not self._has_images(para) and not self._has_math(para):
                continue
            # 첫 유효 문단의 정렬을 셀 정렬로 사용
            if align is None and para.alignment is not None:
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    align = 'center'
                elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    align = 'right'
                elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    align = 'justify'
            has_hyperlink = para._element.find(qn('w:hyperlink')) is not None
            has_math_p = self._has_math(para)
            if has_math_p or has_hyperlink:
                inner = self._process_paragraph_children(para, image_map, options)
            else:
                inner = self._process_runs(para.runs, image_map, options)
            if inner.strip():
                parts.append(inner)
        content = '<br>'.join(parts) if parts else ''
        return content, align

    def _find_tc_at_grid_col(self, tr_elem, target_col):
        """<w:tr> 내에서 grid column 위치에 해당하는 <w:tc> 요소 반환"""
        col = 0
        for tc in tr_elem.findall(qn('w:tc')):
            if col == target_col:
                return tc
            tcPr = tc.find(qn('w:tcPr'))
            span = 1
            if tcPr is not None:
                gs = tcPr.find(qn('w:gridSpan'))
                if gs is not None:
                    try:
                        span = int(gs.get(qn('w:val'), '1'))
                    except ValueError:
                        pass
            col += span
            if col > target_col:
                break
        return None

    def _calc_rowspan(self, tr_elements, start_row, grid_col):
        """start_row부터 아래로 vMerge continuation 셀 개수 계산 (raw XML 기반)"""
        rowspan = 1
        for r in range(start_row + 1, len(tr_elements)):
            tc = self._find_tc_at_grid_col(tr_elements[r], grid_col)
            if tc is None:
                break
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                break
            vm = tcPr.find(qn('w:vMerge'))
            if vm is None:
                break
            if vm.get(qn('w:val')) == 'restart':
                break
            rowspan += 1
        return rowspan

    def _process_table(self, table, image_map, options):
        """표를 HTML table로 변환 (raw XML 기반 병합 + 리치 포매팅)"""
        from docx.table import _Cell

        tbl_elem = table._tbl
        tr_elements = tbl_elem.findall(qn('w:tr'))
        rows_html = []
        skip_cells = set()  # (row_idx, grid_col) — vMerge continuation

        # 멀티행 헤더 감지: Row 0의 vMerge restart가 걸리는 최대 rowspan
        header_rows = 1
        first_tr = tr_elements[0] if tr_elements else None
        if first_tr is not None:
            for tc in first_tr.findall(qn('w:tc')):
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    vm = tcPr.find(qn('w:vMerge'))
                    if vm is not None and vm.get(qn('w:val')) == 'restart':
                        gc = 0
                        # 이 tc의 grid column 계산
                        for prev_tc in first_tr.findall(qn('w:tc')):
                            if prev_tc is tc:
                                break
                            p = prev_tc.find(qn('w:tcPr'))
                            s = 1
                            if p is not None:
                                gs = p.find(qn('w:gridSpan'))
                                if gs is not None:
                                    try:
                                        s = int(gs.get(qn('w:val'), '1'))
                                    except ValueError:
                                        pass
                            gc += s
                        rs = self._calc_rowspan(tr_elements, 0, gc)
                        header_rows = max(header_rows, rs)

        for i, tr in enumerate(tr_elements):
            cells_html = []
            grid_col = 0

            for tc in tr.findall(qn('w:tc')):
                tcPr = tc.find(qn('w:tcPr'))

                # colspan (gridSpan)
                colspan = 1
                if tcPr is not None:
                    gs = tcPr.find(qn('w:gridSpan'))
                    if gs is not None:
                        try:
                            colspan = int(gs.get(qn('w:val'), '1'))
                        except ValueError:
                            pass

                # vMerge continuation → 스킵
                if (i, grid_col) in skip_cells:
                    grid_col += colspan
                    continue

                # vMerge 처리
                rowspan = 1
                if tcPr is not None:
                    vm = tcPr.find(qn('w:vMerge'))
                    if vm is not None:
                        vm_val = vm.get(qn('w:val'))
                        if vm_val == 'restart':
                            rowspan = self._calc_rowspan(tr_elements, i, grid_col)
                            for r in range(i + 1, i + rowspan):
                                for c in range(grid_col, grid_col + colspan):
                                    skip_cells.add((r, c))
                        else:
                            # continuation — 스킵
                            grid_col += colspan
                            continue

                # 셀 내용 (서식 보존)
                cell = _Cell(tc, table)
                content, cell_align = self._process_cell_content(cell, image_map, options)

                # 속성 구성
                attrs = ''
                if colspan > 1:
                    attrs += f' colspan="{colspan}"'
                if rowspan > 1:
                    attrs += f' rowspan="{rowspan}"'
                if cell_align:
                    attrs += f' style="text-align: {cell_align}"'

                tag = 'th' if i < header_rows else 'td'
                cells_html.append(f'<{tag}{attrs}>{content}</{tag}>')
                grid_col += colspan

            if cells_html:
                rows_html.append('<tr>' + ''.join(cells_html) + '</tr>')

        # thead/tbody 분리
        if len(rows_html) > 1 and header_rows > 0:
            thead = '<thead>' + ''.join(rows_html[:header_rows]) + '</thead>'
            tbody = '<tbody>' + ''.join(rows_html[header_rows:]) + '</tbody>'
            return f'<table>{thead}{tbody}</table>'
        elif rows_html:
            return f'<table><tbody>{"".join(rows_html)}</tbody></table>'
        else:
            return ''

    def _process_images(self, doc, image_dir, output_path):
        """
        문서의 이미지를 추출하고 저장

        Args:
            doc: Document 객체
            image_dir: 이미지 저장 디렉토리
            output_path: HTML 출력 경로 (상대 경로 계산용)

        Returns:
            dict: {rId: 상대경로} 매핑
        """
        image_map = {}

        try:
            # 이미지 디렉토리 생성
            ensure_dir(image_dir)

            # 문서의 모든 관계(relationship)에서 이미지 추출
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.reltype:
                    try:
                        image_data = rel.target_part.blob
                        content_type = rel.target_part.content_type

                        # 확장자 결정
                        ext_map = {
                            'image/png': '.png',
                            'image/jpeg': '.jpg',
                            'image/gif': '.gif',
                            'image/bmp': '.bmp',
                            'image/tiff': '.tiff',
                            'image/x-emf': '.emf',
                            'image/x-wmf': '.wmf',
                        }
                        ext = ext_map.get(content_type, '.png')

                        # 파일명 생성 (해시 기반)
                        hash_name = hashlib.md5(image_data).hexdigest()[:12]
                        filename = f"image_{hash_name}{ext}"
                        image_path = image_dir / filename

                        # 이미지 저장
                        with open(image_path, 'wb') as f:
                            f.write(image_data)

                        # 경로 계산: image_prefix 지정 시 접두사+파일명, 아니면 상대경로 자동
                        if getattr(self, '_image_prefix', None):
                            prefix = self._image_prefix.rstrip('/')
                            rel_path = f"{prefix}/{filename}"
                        else:
                            rel_path = os.path.relpath(image_path, output_path.parent)
                            rel_path = rel_path.replace('\\', '/')

                        image_map[rel_id] = rel_path
                        self.logger.debug(f"이미지 추출: {filename}")

                    except Exception as e:
                        self.logger.warning(f"이미지 추출 실패 ({rel_id}): {e}")

        except Exception as e:
            self.logger.error(f"이미지 처리 중 오류: {e}")

        return image_map

    def _has_images(self, paragraph):
        """문단에 이미지가 있는지 확인"""
        # XML에서 drawing 요소 확인
        drawing_ns = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
        for elem in paragraph._element.iter():
            if 'drawing' in elem.tag or 'pict' in elem.tag:
                return True
        return False

    def _has_unextractable_shapes(self, paragraph):
        """문단에 이미지로 추출 불가능한 Word 도형/그리기가 있는지 확인.

        DrawingML(<w:drawing>)이나 VML(<w:pict>)이 존재하지만
        blip(임베디드 이미지)이 전혀 없는 경우 True 반환.
        """
        p_elem = paragraph._element
        has_drawing = False
        has_blip = False
        for elem in p_elem.iter():
            tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
            if tag in ('drawing', 'pict'):
                has_drawing = True
            if tag == 'blip':
                has_blip = True
        return has_drawing and not has_blip

    def _run_has_shape(self, run):
        """Run에 drawing/pict 요소가 있는지 확인 (blip 없이)"""
        for elem in run._element.iter():
            tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
            if tag in ('drawing', 'pict'):
                return True
        return False

    def _has_math(self, paragraph):
        """문단에 수식(OMML)이 있는지 확인"""
        p_elem = paragraph._element
        return (p_elem.find(qn('m:oMathPara')) is not None or
                p_elem.find(qn('m:oMath')) is not None)

    def _extract_footnotes(self, doc):
        """문서에서 각주 추출 → {id: text} 맵 반환"""
        footnotes = {}
        try:
            for rel in doc.part.rels.values():
                if 'footnotes' in rel.reltype:
                    fn_xml = ET.fromstring(rel.target_part.blob)
                    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                    for fn in fn_xml.findall(f'{{{ns}}}footnote'):
                        fn_id = fn.get(f'{{{ns}}}id')
                        if fn_id in ('0', '-1'):
                            continue
                        texts = []
                        for t in fn.iter(f'{{{ns}}}t'):
                            if t.text:
                                texts.append(t.text)
                        if texts:
                            footnotes[fn_id] = ''.join(texts)
                    break
        except Exception as e:
            self.logger.debug(f"각주 추출 실패: {e}")
        return footnotes

    def _extract_endnotes(self, doc):
        """문서에서 미주 추출 → {id: text} 맵 반환"""
        endnotes = {}
        try:
            for rel in doc.part.rels.values():
                if 'endnotes' in rel.reltype:
                    en_xml = ET.fromstring(rel.target_part.blob)
                    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                    for en in en_xml.findall(f'{{{ns}}}endnote'):
                        en_id = en.get(f'{{{ns}}}id')
                        if en_id in ('0', '-1'):
                            continue
                        texts = []
                        for t in en.iter(f'{{{ns}}}t'):
                            if t.text:
                                texts.append(t.text)
                        if texts:
                            endnotes[en_id] = ''.join(texts)
                    break
        except Exception as e:
            self.logger.debug(f"미주 추출 실패: {e}")
        return endnotes

    def _get_image_width_emu(self, blip_elem):
        """blip 요소에서 Word 배치 폭(EMU) 추출"""
        parent = blip_elem
        for _ in range(10):
            parent = parent.getparent()
            if parent is None:
                break
            ptag = parent.tag.split('}')[-1] if '}' in parent.tag else parent.tag
            if ptag in ('inline', 'anchor'):
                for child in parent:
                    ctag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if ctag == 'extent':
                        cx = child.get('cx')
                        if cx:
                            return int(cx)
                break
        return None

    def _get_page_content_width(self, doc):
        """문서의 페이지 콘텐츠 폭(EMU) 계산 — 페이지 폭 - 좌우 여백"""
        try:
            body = doc.element.find(qn('w:body'))
            sectPr = body.find(qn('w:sectPr'))
            pgSz = sectPr.find(qn('w:pgSz'))
            pgMar = sectPr.find(qn('w:pgMar'))
            page_w = int(pgSz.get(qn('w:w')))
            margin_l = int(pgMar.get(qn('w:left')))
            margin_r = int(pgMar.get(qn('w:right')))
            return (page_w - margin_l - margin_r) * 635  # twips → EMU
        except Exception:
            return 5731510  # A4 기본값 (210mm - 1in*2 여백)

    def _make_img_tag(self, src, blip_elem=None):
        """<img> 태그 생성 — Word 페이지 대비 비율(%)로 출력"""
        attrs = f'src="{src}" alt="" decoding="async"'
        if blip_elem is not None:
            cx = self._get_image_width_emu(blip_elem)
            if cx and self._page_content_width:
                pct = round(cx / self._page_content_width * 100)
                pct = min(pct, 100)  # 100% 상한
                attrs += f' style="width: {pct}%"'
        return f'<img {attrs}>'

    def _get_run_image(self, run):
        """Run에서 이미지 (rId, blip_elem) 추출"""
        try:
            for elem in run._element.iter():
                if 'blip' in elem.tag:
                    embed = elem.get(qn('r:embed'))
                    if embed:
                        return embed, elem
        except:
            pass
        return None, None

    def _extract_inline_images(self, paragraph, image_map, align_attr=''):
        """문단에서 인라인 이미지 추출"""
        images = []

        try:
            for elem in paragraph._element.iter():
                if 'blip' in elem.tag:
                    embed = elem.get(qn('r:embed'))
                    if embed and embed in image_map:
                        images.append(self._make_img_tag(image_map[embed], elem))
        except:
            pass

        if images:
            return f'<p{align_attr}>' + ''.join(images) + '</p>'
        return None

    def _detect_special_block(self, text):
        """
        특수 블록 (NOTE, WARNING, CAUTION) 감지

        Args:
            text: 문단 텍스트

        Returns:
            tuple: (블록 타입, 내용) 또는 None
        """
        special_blocks = self.config.get('special_blocks', {})

        for block_type, keywords in special_blocks.items():
            for keyword in keywords:
                # "NOTE: 내용" 또는 "NOTE - 내용" 패턴
                patterns = [
                    rf'^{re.escape(keyword)}\s*[:：\-]\s*(.+)$',
                    rf'^【{re.escape(keyword)}】\s*(.+)$',
                    rf'^\[{re.escape(keyword)}\]\s*(.+)$',
                ]

                for pattern in patterns:
                    match = re.match(pattern, text, re.IGNORECASE)
                    if match:
                        return (block_type, match.group(1).strip())

        return None

    # ── 캡션 자동 ID 부여 + 참조 하이퍼링크 ──────────────────────────

    def _resolve_styleref_fields(self, paragraph):
        """STYLEREF 필드 해석 (Plan-37 Phase 4b).

        `{ STYLEREF "Heading 1" [\\s] [\\n] [\\w] }` 같은 필드를 _heading_context
        에서 해당 level 의 최신 번호로 치환. 캡션의 "그림 3-1" 합성 지원.

        지원:
          - STYLEREF "Heading N" 또는 STYLEREF N → level N heading 의 번호
          - STYLEREF "제목 N" (한글)
          - 스위치 (\\s, \\n, \\w 등) 는 인식만 하고 단순 번호 치환
        """
        if not hasattr(self, '_heading_context'):
            return
        p_elem = paragraph._element

        # 1. 단순 필드 (fldSimple)
        for fld_simple in p_elem.findall(qn('w:fldSimple')):
            instr = fld_simple.get(qn('w:instr'), '')
            level = self._parse_styleref_level(instr)
            if level is None:
                continue
            value = self._heading_context.get(level, "")
            if value:
                child_runs = fld_simple.findall(qn('w:r'))
                self._set_field_text(child_runs, value, fld_simple)

        # 2. 복합 필드 (fldChar)
        state = 'normal'
        level = None
        instr_text = ''
        result_runs = []
        for child in list(p_elem):
            if child.tag != qn('w:r'):
                continue
            fld_char = child.find(qn('w:fldChar'))
            if fld_char is not None:
                ft = fld_char.get(qn('w:fldCharType'))
                if ft == 'begin':
                    state, level, instr_text, result_runs = 'in_field', None, '', []
                elif ft == 'separate':
                    if state == 'in_field':
                        level = self._parse_styleref_level(instr_text)
                    state = 'in_result' if level is not None else 'normal'
                elif ft == 'end':
                    if state == 'in_result' and level is not None:
                        value = self._heading_context.get(level, "")
                        if value:
                            self._set_field_text(result_runs, value)
                    state, level, result_runs = 'normal', None, []
                continue
            if state == 'in_field':
                instr_el = child.find(qn('w:instrText'))
                if instr_el is not None and instr_el.text:
                    instr_text += instr_el.text
            elif state == 'in_result':
                result_runs.append(child)

    def _parse_styleref_level(self, instr):
        """STYLEREF 필드 instr 에서 level 정수 추출.

        예: 'STYLEREF "Heading 1" \\s' → 1, 'STYLEREF 2' → 2
        """
        if 'STYLEREF' not in instr.upper():
            return None
        # "Heading N" / "제목 N" / 단순 숫자
        m = re.search(r'(?:Heading|제목)\s*(\d+)', instr, re.IGNORECASE)
        if m:
            return int(m.group(1))
        m = re.search(r'STYLEREF\s+"?(\d+)"?', instr, re.IGNORECASE)
        if m:
            return int(m.group(1))
        return None

    def _parse_seq_instr(self, instr):
        """SEQ 필드 instr 텍스트 파싱 (Plan-37 Phase 4c).

        문법:  SEQ {category} [\\r N] [\\s N] [\\c] [\\n] [\\* FORMAT] [\\h]

        Returns:
            dict 또는 None — {category, reset_to, heading_reset, repeat, format}
        """
        m = re.search(r'SEQ\s+(\S+)', instr, re.IGNORECASE)
        if not m:
            return None
        info = {
            'category': m.group(1),
            'reset_to': None,       # \r N
            'heading_reset': None,  # \s N (level)
            'repeat': False,        # \c
            'format': 'decimal',    # \* ARABIC|ROMAN|alphabetic...
        }
        # \r N
        m = re.search(r'\\r\s*(\d+)', instr, re.IGNORECASE)
        if m:
            info['reset_to'] = int(m.group(1))
        # \s N
        m = re.search(r'\\s\s*(\d+)', instr, re.IGNORECASE)
        if m:
            info['heading_reset'] = int(m.group(1))
        # \c
        if re.search(r'\\c\b', instr, re.IGNORECASE):
            info['repeat'] = True
        # \* FORMAT
        m = re.search(r'\\\*\s*(\w+)', instr, re.IGNORECASE)
        if m:
            fmt = m.group(1).upper()
            info['format'] = {
                'ARABIC': 'decimal',
                'ARABICDASH': 'decimal',
                'ROMAN': 'upperRoman',
                'roman': 'lowerRoman',
                'ALPHABETIC': 'upperLetter',
                'alphabetic': 'lowerLetter',
            }.get(fmt, 'decimal')
        return info

    def _compute_seq_value(self, info):
        """SEQ 필드 정보로 다음 값 계산 + counter 갱신.

        heading_reset: heading level 기준 리셋 — 이 level 번호가 바뀔 때마다 리셋.
        """
        category = info['category']
        # \s N: heading level N 변화 감지
        if info['heading_reset'] is not None:
            h_key = f'_hres_{category}_{info["heading_reset"]}'
            current_h = self._heading_context.get(info['heading_reset'], 0)
            last_h = self._seq_counters.get(h_key, None)
            if last_h != current_h:
                # heading 변경 → counter 리셋
                self._seq_counters[category] = 0
                self._seq_counters[h_key] = current_h

        if info['repeat']:
            # \c — 현재 값 그대로 반환 (증가 X)
            return self._seq_counters.get(category, 1)

        if info['reset_to'] is not None:
            # \r N — 해당 값으로 설정
            self._seq_counters[category] = info['reset_to']
            return info['reset_to']

        # 기본 \n — 증가
        self._seq_counters[category] = self._seq_counters.get(category, 0) + 1
        return self._seq_counters[category]

    def _resolve_seq_fields(self, paragraph):
        """문단 XML에서 SEQ 필드를 감지하고 빈 캐시 값을 자동 채번.

        복합 필드(fldChar begin/separate/end)와 단순 필드(fldSimple)를 모두 처리.
        Phase 4c: `\\r`, `\\s`, `\\c`, `\\*` 스위치 지원.
        """
        from numbering_resolver import _format_number as _fmt_num

        p_elem = paragraph._element

        # 1. 단순 필드(fldSimple) → 자식 run을 부모 <w:p>로 승격
        for fld_simple in p_elem.findall(qn('w:fldSimple')):
            instr = fld_simple.get(qn('w:instr'), '')
            info = self._parse_seq_instr(instr)
            if not info:
                continue

            category = info['category']
            child_runs = fld_simple.findall(qn('w:r'))
            cached = self._get_field_text(child_runs)

            if cached and cached.isdigit() and not info['repeat'] \
                    and info['reset_to'] is None and info['heading_reset'] is None:
                # cached 유효 + 스위치 없으면 기존 동작 (cached 존중)
                self._seq_counters[category] = max(
                    self._seq_counters.get(category, 0), int(cached))
            else:
                value = self._compute_seq_value(info)
                formatted = _fmt_num(value, info['format'])
                self._set_field_text(child_runs, formatted, fld_simple)

            # fldSimple 자식 run을 <w:p>로 승격
            idx = list(p_elem).index(fld_simple)
            for child in list(child_runs):
                fld_simple.remove(child)
                p_elem.insert(idx, child)
                idx += 1
            p_elem.remove(fld_simple)

        # 2. 복합 필드(fldChar begin → instrText → separate → result → end)
        state = 'normal'
        seq_info = None
        instr_text = ''
        result_runs = []

        for child in list(p_elem):
            if child.tag != qn('w:r'):
                continue

            fld_char = child.find(qn('w:fldChar'))
            if fld_char is not None:
                fld_type = fld_char.get(qn('w:fldCharType'))
                if fld_type == 'begin':
                    state = 'in_field'
                    seq_info = None
                    instr_text = ''
                    result_runs = []
                elif fld_type == 'separate':
                    if state == 'in_field':
                        seq_info = self._parse_seq_instr(instr_text)
                    state = 'in_result' if seq_info else 'normal'
                elif fld_type == 'end':
                    if state == 'in_result' and seq_info:
                        category = seq_info['category']
                        cached = self._get_field_text(result_runs)
                        if cached and cached.isdigit() and not seq_info['repeat'] \
                                and seq_info['reset_to'] is None \
                                and seq_info['heading_reset'] is None:
                            self._seq_counters[category] = max(
                                self._seq_counters.get(category, 0),
                                int(cached))
                        else:
                            value = self._compute_seq_value(seq_info)
                            formatted = _fmt_num(value, seq_info['format'])
                            self._set_field_text(result_runs, formatted)
                    state = 'normal'
                    seq_info = None
                    result_runs = []
                continue

            if state == 'in_field':
                instr_el = child.find(qn('w:instrText'))
                if instr_el is not None and instr_el.text:
                    instr_text += instr_el.text
            elif state == 'in_result':
                result_runs.append(child)

    def _get_field_text(self, runs):
        """run 요소 리스트에서 <w:t> 텍스트를 결합하여 반환"""
        parts = []
        for r in runs:
            t = r.find(qn('w:t'))
            if t is not None and t.text:
                parts.append(t.text)
        return ''.join(parts).strip()

    def _set_field_text(self, runs, text, parent_for_new=None):
        """run 리스트의 첫 번째 run에 텍스트를 설정 (없으면 생성)"""
        if runs:
            t = runs[0].find(qn('w:t'))
            if t is None:
                t = runs[0].makeelement(qn('w:t'), {})
                runs[0].append(t)
            t.text = text
        elif parent_for_new is not None:
            # run이 없으면 새로 생성
            r = parent_for_new.makeelement(qn('w:r'), {})
            t = r.makeelement(qn('w:t'), {})
            t.text = text
            r.append(t)
            parent_for_new.append(r)

    def _is_display_caption(self, text):
        """표시 캡션 판정 — class="caption" 만 부여할지 (Plan-73).

        _detect_caption(참조 캡션) 이 놓치는 표기까지 넓게 잡는다.
        구분자 없는 "표 1 시스템 구성", 붙여쓴 "표1", "Table 1 Overview" 등
        담당자별 표기 편차가 여기서 흡수된다.

        id 나 본문 참조 링크는 만들지 않는다 — 판정이 느슨한 만큼 오탐이
        섞이므로, 오탐의 피해를 "가운데 정렬 + 마진"으로 한정한다.

        본문 오탐 방어선 2겹:
          1. 조사 배제 — "표 1을 보면", "그림 3과 같이" 처럼 번호에 조사가
             바로 붙으면 캡션이 아니라 서술이다. 짧아도 걸러진다.
          2. 길이 가드 — 조사 없이 길게 이어지는 설명 문장을 거른다.
        """
        return (len(text) < DISPLAY_CAPTION_MAX_LEN
                and bool(DISPLAY_CAPTION_RE.match(text))
                and not DISPLAY_CAPTION_PARTICLE_RE.match(text))

    def _detect_caption(self, text, paragraph=None):
        """참조 캡션 감지 → ID 생성 및 _caption_map 등록 (Plan-37 Phase 4d).

        여기서 감지된 캡션만 id 를 받고 본문 참조 링크(_linkify_references)의
        대상이 된다. 표시 전용 판정은 _is_display_caption 참조 (Plan-73).

        감지 규칙:
          1. 기본: "Figure N: ..." 또는 "Figure N. ..." 등 구분자 포함
          2. 확장(Phase 4d): 구분자 없이도 캡션 스타일(Caption/캡션/제목 주석)로
             감지된 단락은 캡션으로 인정
          3. 구분자 후보: `:` `：` `–` `—` `-` `.` `]` `>`
          4. 제목 문맥 ("Figure 1" 으로 시작하는 단락 전체가 짧고 제목 형태) 허용

        "그림 1 또한 중요하다" 같은 본문 오탐을 피하기 위해 기본은 구분자 필수.
        캡션 스타일이 명시된 단락은 구분자 없이도 감지.

        Args:
            text: 단락 텍스트
            paragraph: python-docx Paragraph 객체 (style 확인용, 선택)

        Returns:
            str or None: 캡션 ID (e.g. 'fig-1', 'tbl-2-1')
        """
        # 1차: 기존 엄격 규칙 (구분자 필수)
        strict_match = re.match(
            r'^(?:Figure|Fig\.?|Table|Tab\.?|그림|표)\s+'
            r'\d+(?:[-.]?\d+)*'
            r'\s*[:：–—\-.\]\>]',
            text, re.IGNORECASE)

        is_caption = bool(strict_match)

        # 2차: 캡션 스타일 직접 감지 (스타일 ID 또는 이름에 'caption' 또는 '캡션' 포함)
        if not is_caption and paragraph is not None and paragraph.style:
            sid = (paragraph.style.style_id or "").lower()
            sname = (paragraph.style.name or "").lower()
            if 'caption' in sid or 'caption' in sname or '캡션' in sname:
                # 캡션 스타일이면 "Figure N" 숫자만 있어도 허용
                relaxed = re.match(
                    r'^(?:Figure|Fig\.?|Table|Tab\.?|그림|표)\s+'
                    r'\d+(?:[-.]?\d+)*',
                    text, re.IGNORECASE)
                is_caption = bool(relaxed)

        if not is_caption:
            return None

        caption_id = self._make_caption_id(text)
        if caption_id:
            self._caption_map[caption_id] = text
        return caption_id

    def _make_caption_id(self, text):
        """캡션/참조 텍스트에서 ID 생성 (e.g. 'Figure 1' → 'fig-1')"""
        m = re.match(
            r'^(Figure|Fig\.?|Table|Tab\.?|그림|표)\s+(\d+(?:[-.]?\d+)*)',
            text, re.IGNORECASE)
        if not m:
            return None

        keyword = m.group(1).rstrip('.').lower()
        number = m.group(2).replace('.', '-')

        prefix_map = {
            'figure': 'fig', 'fig': 'fig',
            'table': 'tbl', 'tab': 'tbl',
            '그림': 'fig', '표': 'tbl',
        }
        prefix = prefix_map.get(keyword, 'fig')
        return f'{prefix}-{number}'

    def _embed_provenance(self, html_content, options):
        """출력 HTML 에 변환기 버전·어댑터·일시를 embed (Plan-37 Phase 3f).

        fragment 모드 (output.fragment_only=True): HTML 주석으로 삽입
        전체 HTML 모드: <meta name="..."> 태그 삽입

        adapter 이름은 호출 컨텍스트에 없으므로 선택적으로 self._provenance_adapter
        속성을 사용. 없으면 'unknown' 으로 기록.
        """
        try:
            from __version__ import __version__ as _ver
        except ImportError:
            _ver = "unknown"

        from datetime import datetime, timezone
        now = datetime.now(timezone.utc).replace(microsecond=0).isoformat()
        adapter = getattr(self, '_provenance_adapter', None) or 'unknown'

        fragment_only = options.get('fragment_only',
                                    self.config.get('output', {}).get('fragment_only', True))

        if fragment_only:
            # HTML 주석 — Explorer 가 fragment 로 주입하므로 주석 형식
            banner = (
                f'<!-- converter: smart-doc-platform/docx-converter {_ver} '
                f'| adapter: {adapter} | date: {now} -->'
            )
            return banner + '\n' + html_content
        else:
            # 전체 HTML 모드 — <head> 에 meta 삽입
            meta = (
                f'<meta name="converter" content="smart-doc-platform/docx-converter">\n'
                f'<meta name="converter-version" content="{_ver}">\n'
                f'<meta name="converter-adapter" content="{adapter}">\n'
                f'<meta name="conversion-date" content="{now}">'
            )
            if '</head>' in html_content:
                return html_content.replace('</head>', meta + '\n</head>', 1)
            return meta + '\n' + html_content

    def _linkify_references(self, html_content):
        """최종 HTML에서 캡션 참조 텍스트를 <a data-fig-ref> 링크로 변환.

        - _caption_map에 등록된 캡션만 링크 생성
        - 캡션 요소 자체(id="fig-/tbl-"), 기존 <a> 태그 내부는 스킵
        """
        if not self._caption_map:
            return html_content

        ref_kw = r'(?:Figure|Fig\.?|Table|Tab\.?|그림|표)'
        ref_num = r'\d+[-.]?\d*'
        ref_suffix = r'(?:\s*(?:참조|참고))?'

        combined = re.compile(
            rf'(<a\s[^>]*>.*?</a>)'          # (1) 기존 <a> 태그 → 스킵
            rf'|(<[^>]+>)'                     # (2) HTML 태그 → 스킵
            rf'|({ref_kw}\s+{ref_num}{ref_suffix})',  # (3) 참조 텍스트
            re.IGNORECASE | re.DOTALL)

        lines = html_content.split('\n')
        result = []

        for line in lines:
            # 캡션 요소 자체는 스킵
            if re.search(r'\bid="(?:fig|tbl)-', line):
                result.append(line)
                continue

            def _replacer(m):
                if m.group(1) or m.group(2):
                    return m.group(0)
                ref_text = m.group(3)
                cap_id = self._make_caption_id(ref_text)
                if cap_id and cap_id in self._caption_map:
                    return f'<a data-fig-ref="{cap_id}">{ref_text}</a>'
                return ref_text

            result.append(combined.sub(_replacer, line))

        return '\n'.join(result)

    def analyze(self, input_path):
        """
        문서 구조 분석 (미리보기용)

        Args:
            input_path: 입력 .docx 파일 경로

        Returns:
            dict: 문서 구조 정보
        """
        input_path = Path(input_path)

        if not input_path.exists():
            return {'error': f"파일을 찾을 수 없습니다: {input_path}"}

        try:
            doc = Document(str(input_path))

            # 기본 정보
            analysis = {
                'filename': input_path.name,
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'sections': len(doc.sections),
                'headings': {'h1': 0, 'h2': 0, 'h3': 0, 'h4': 0, 'h5': 0, 'h6': 0},
                'images': 0,
                'styles_used': set(),
                'font_sizes_used': set(),
                'starts_with_h1': False,
                'warnings': []
            }

            # 이미지 수 계산
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.reltype:
                    analysis['images'] += 1

            # 문단 분석
            first_content_found = False
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # 스타일 수집
                if para.style and para.style.name:
                    analysis['styles_used'].add(para.style.name)

                # 폰트 크기 수집
                font_size = self._get_paragraph_font_size(para)
                if font_size:
                    analysis['font_sizes_used'].add(font_size)

                # 제목 레벨 분석
                tag = self._detect_heading_level(para)
                if re.match(r'^h[1-6]$', tag):
                    analysis['headings'][tag] = analysis['headings'].get(tag, 0) + 1

                    if not first_content_found:
                        if tag == 'h1':
                            analysis['starts_with_h1'] = True
                        first_content_found = True
                elif not first_content_found:
                    first_content_found = True

            # 경고 생성
            if not analysis['starts_with_h1']:
                analysis['warnings'].append("문서가 h1으로 시작하지 않습니다.")

            if analysis['headings']['h1'] == 0:
                analysis['warnings'].append("h1 제목이 없습니다.")

            if analysis['headings']['h1'] > 1:
                analysis['warnings'].append(f"h1 제목이 {analysis['headings']['h1']}개 있습니다. (권장: 1개)")

            # set을 list로 변환 (JSON 직렬화용)
            analysis['styles_used'] = sorted(list(analysis['styles_used']))
            analysis['font_sizes_used'] = sorted(list(analysis['font_sizes_used']))

            return analysis

        except Exception as e:
            return {'error': str(e)}


# 테스트용
if __name__ == "__main__":
    converter = DocxConverter()
    print("Config loaded:", json.dumps(converter.config, indent=2, ensure_ascii=False))
