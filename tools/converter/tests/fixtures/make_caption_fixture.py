# -*- coding: utf-8 -*-
"""
Plan-73 — 캡션 표기 편차 검증용 DOCX 생성 스크립트

담당자마다 다른 캡션 표기를 한 문서에 모아, 참조 캡션(id 부여)과
표시 캡션(class 만)의 경계가 의도대로 갈리는지 검증한다.

출력: caption_variants.docx (test_caption_tiers.py 가 없으면 자동 생성)
"""
from pathlib import Path

from docx import Document

# (본문 텍스트, 기대 판정) — 기대 판정은 test_caption_tiers.py 의 CASES 와 함께 관리
PARAGRAPHS = [
    # 참조 캡션 — 구분자 있음 → class + id
    "표 1. 시스템 구성",
    "표 2: 주요 제원",
    "그림 3-1. 흐름도",
    "Table 4. Overview",
    "Tab. 5: Legacy abbreviation",

    # 표시 캡션 — 구분자 없음 → class 만
    "표 6 시스템 구성",
    "표7 붙여쓴 표기",
    "Table 8 Overview",
    "그림 9 흐름도",
    "Fig. 10 Diagram",

    # 캡션 아님 — 조사가 붙은 본문 서술 (짧아도 걸러져야 함)
    "표 11을 보면 알 수 있듯이 구성은 다음과 같다.",
    "표 12는 주요 제원을 정리한 것이다.",
    "그림 13과 같이 배치한다.",
    "그림 14의 흐름을 따른다.",
    "Table 15를 참조한다.",

    # 캡션 아님 — 그 외 본문 오탐 후보
    "그림 16 또한 중요하다는 점을 여기서 길게 설명한다. " * 6,   # 185자 — 길이 가드 초과
    "본 절에서는 시스템 구성을 설명한다.",
    "표 없이 시작하는 문장",
]


def build(output_path):
    doc = Document()
    doc.add_heading("캡션 표기 편차 픽스처", level=1)
    for text in PARAGRAPHS:
        doc.add_paragraph(text)
    doc.save(str(output_path))
    return output_path


if __name__ == "__main__":
    out = Path(__file__).resolve().parent / "caption_variants.docx"
    build(out)
    print(f"생성 완료: {out}")
