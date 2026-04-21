# -*- coding: utf-8 -*-
"""
Plan-37 Phase 4 — STYLEREF + SEQ 스위치 검증용 DOCX 생성 스크립트

python-docx 로 fixture 생성:
  - heading 1 여러 개 + 각 chapter 안에 Figure 캡션
  - "그림 {STYLEREF 1 \\s}-{SEQ Figure \\s 1}" 패턴 (heading 1 변경 시 SEQ 리셋)
  - SEQ 스위치: \\r 10 (10부터 시작), \\c (repeat)

출력: phase4_styleref.docx (테스트 시 동적 생성 가능)
"""
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _add_fld(paragraph, instr, cached_value=""):
    """fldChar begin/separate/end 형태로 필드 추가."""
    r1 = paragraph.add_run()
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    r1._r.append(fc1)

    r2 = paragraph.add_run()
    it = OxmlElement('w:instrText')
    it.text = instr
    r2._r.append(it)

    r3 = paragraph.add_run()
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'separate')
    r3._r.append(fc2)

    r4 = paragraph.add_run(cached_value)

    r5 = paragraph.add_run()
    fc3 = OxmlElement('w:fldChar')
    fc3.set(qn('w:fldCharType'), 'end')
    r5._r.append(fc3)


def build(output_path):
    doc = Document()
    # Chapter 1
    doc.add_heading("Chapter One", level=1)
    doc.add_paragraph("본문 내용.")
    # 캡션 1: "그림 1-1: ..." (STYLEREF heading1 + SEQ Figure \s 1)
    p = doc.add_paragraph()
    p.add_run("그림 ")
    _add_fld(p, 'STYLEREF 1 \\s', '')  # 현재 heading 1 번호
    p.add_run("-")
    _add_fld(p, 'SEQ Figure \\s 1', '')  # SEQ Figure, heading 1 변경 시 리셋
    p.add_run(": 첫 번째 그림")

    # 캡션 2: 같은 chapter → Figure 2
    p = doc.add_paragraph()
    p.add_run("그림 ")
    _add_fld(p, 'STYLEREF 1 \\s', '')
    p.add_run("-")
    _add_fld(p, 'SEQ Figure \\s 1', '')
    p.add_run(": 두 번째 그림")

    # Chapter 2 → SEQ Figure 리셋 (1부터 다시)
    doc.add_heading("Chapter Two", level=1)
    doc.add_paragraph("두 번째 챕터.")
    p = doc.add_paragraph()
    p.add_run("그림 ")
    _add_fld(p, 'STYLEREF 1 \\s', '')
    p.add_run("-")
    _add_fld(p, 'SEQ Figure \\s 1', '')
    p.add_run(": 챕터 2 의 첫 그림")

    # SEQ \\r 10 — 10부터 시작
    p = doc.add_paragraph()
    p.add_run("Table ")
    _add_fld(p, 'SEQ Table \\r 10', '')
    p.add_run(": 10 번 표")

    # SEQ \\c — 반복
    p = doc.add_paragraph()
    p.add_run("참조: Table ")
    _add_fld(p, 'SEQ Table \\c', '')
    p.add_run(" 참고")

    doc.save(str(output_path))
    print(f"fixture 생성: {output_path}")


if __name__ == "__main__":
    out = Path(__file__).parent / "phase4_styleref.docx"
    build(out)
